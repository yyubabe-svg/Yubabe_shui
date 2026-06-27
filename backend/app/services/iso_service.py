import os
import re
import json
import uuid
import shutil
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Pt

from app.core.config import settings
from app.schemas.iso import ISOProjectInfo
from app.services.llm_service import llm_service


ALL_MAJORS = [
    "工程测量", "工程地质", "水文", "规划/节水", "水工建筑物",
    "土建/管理/安全", "水机/暖通", "金属结构", "电气一次", "电气二次",
    "信息化", "消防", "施工/节能", "环境保护", "水土保持",
    "造价", "经济评价", "水文化", "其他专业"
]


ISO_EXTRACT_PROMPT = """你是水利工程设计文件信息提取专家，请从以下项目设计报告中提取关键信息，用于填写ISO管理体系附表。

请严格按照JSON格式输出，确保所有字段准确填写。如果某项信息无法从报告中确定，请填null。

需要提取的字段说明：
1. project_name: 工程项目全称
2. project_code: 项目编码（如CDSD260001C），从文件名或封面提取
3. feature_code: 特征码（项目编码最后一位字母）
4. design_stage: 设计阶段，只能是"初步设计"、"可行性研究"、"实施方案"之一
5. report_date: 报告日期（如2026年1月）
6. client: 业主单位/建设单位/委托单位名称
7. work_scope: 工作内容描述（100-200字，包含具体工程量指标）
8. engineering_overview: 工程概况及主要任务（150-250字）
9. design_basis: 设计依据
10. technical_points: 主要技术要点及难点（100-200字）
11. risk_points: 主要风险识别及应对措施（100-200字）
12. customer_requirements: 顾客主要技术要求
13. external_resources: 外部资源要求
14. project_grade: 工程等别（如Ⅳ等、Ⅴ等）
15. building_level: 主要建筑物级别（如4级、5级）
16. flood_standard: 防洪标准（如20年一遇）
17. drainage_standard: 排涝标准（如10年一遇，如无则填null）
18. involved_majors: 涉及的专业列表（从固定列表中选择）
19. applicable_codes: 适用的主要规范列表，格式为[{"name": "规范名称", "code": "编号"}]

固定专业列表：工程测量、工程地质、水文、规划/节水、水工建筑物、土建/管理/安全、水机/暖通、金属结构、电气一次、电气二次、信息化、消防、施工/节能、环境保护、水土保持、造价、经济评价、水文化、其他专业

重要提示：
- 工作内容必须包含具体量化指标
- 小型河道治理、堤防工程、移民后扶项目通常质量分级为B级
- 根据项目类型判断涉及专业
- 可研阶段一般不涉及经济评价

项目报告内容：
{report_content}

请输出JSON："""


class ISOService:
    """ISO管理体系文档自动填写服务"""
    
    def __init__(self):
        self.template_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            "resources", "iso_templates"
        )
        self.template_path = os.path.join(self.template_dir, "管理体系附表-设计部分.docx")
        self.output_dir = os.path.join(settings.UPLOAD_DIR, "iso_output")
        os.makedirs(self.output_dir, exist_ok=True)
        self.tasks: Dict[str, Dict[str, Any]] = {}
    
    def parse_filename(self, filename: str) -> Tuple[Optional[str], Optional[str]]:
        """从文件名解析项目编码和名称"""
        pattern = r'(CDSD\d{2}\d{3}[A-Z])[_]?(.+?)(?:_管理体系|_初设|_可行性|_实施方案|$)'
        match = re.search(pattern, filename)
        if match:
            code = match.group(1)
            name = match.group(2).strip()
            return code, name
        return None, None
    
    def detect_design_stage(self, text: str, filename: str = "") -> str:
        """检测设计阶段"""
        combined = text + filename
        if "初步设计" in combined or "初设" in combined:
            return "初步设计"
        elif "可行性研究" in combined or "可研" in combined:
            return "可行性研究"
        elif "实施方案" in combined:
            return "实施方案"
        return "初步设计"
    
    def determine_quality_level(self, project_info: Dict[str, Any]) -> str:
        """判定质量控制分级"""
        grade = project_info.get("project_grade", "") or ""
        level = project_info.get("building_level", "") or ""
        
        if "Ⅰ等" in grade or "1级" in level:
            return "A级"
        elif "Ⅱ等" in grade or "Ⅲ等" in grade or "2级" in level or "3级" in level:
            return "B级"
        else:
            return "B级"
    
    def determine_majors(self, design_stage: str) -> Tuple[List[str], List[str]]:
        """根据设计阶段确定涉及专业"""
        if design_stage == "可行性研究":
            involved = ["工程测量", "工程地质", "水文", "规划/节水", "水工建筑物", 
                       "土建/管理/安全", "施工/节能", "环境保护", "水土保持", "造价"]
        elif design_stage == "初步设计":
            involved = ["工程测量", "工程地质", "水文", "规划/节水", "水工建筑物",
                       "土建/管理/安全", "金属结构", "电气一次", "消防", "施工/节能",
                       "环境保护", "水土保持", "造价"]
        else:
            involved = ["工程测量", "工程地质", "水文", "规划/节水", "水工建筑物",
                       "土建/管理/安全", "水机/暖通", "金属结构", "消防", "施工/节能",
                       "环境保护", "水土保持", "造价", "经济评价"]
        
        excluded = [m for m in ALL_MAJORS if m not in involved]
        return involved, excluded
    
    async def extract_project_info(self, text: str, filename: str = "") -> ISOProjectInfo:
        """从报告文本中提取项目信息"""
        code_from_name, name_from_name = self.parse_filename(filename)
        design_stage = self.detect_design_stage(text, filename)
        
        truncated_text = text[:8000] if len(text) > 8000 else text
        
        prompt = ISO_EXTRACT_PROMPT.format(report_content=truncated_text)
        
        try:
            llm_response = await llm_service.generate(prompt)
            json_match = re.search(r'\{[\s\S]*\}', llm_response)
            if json_match:
                data = json.loads(json_match.group())
            else:
                data = {}
        except Exception as e:
            print(f"LLM提取失败: {e}")
            data = {}
        
        project_name = data.get("project_name") or name_from_name or "未知项目"
        project_code = data.get("project_code") or code_from_name or "CDSD260000X"
        feature_code = data.get("feature_code") or (project_code[-1] if project_code else "X")
        
        if not data.get("involved_majors"):
            involved, excluded = self.determine_majors(design_stage)
        else:
            involved = data["involved_majors"]
            excluded = [m for m in ALL_MAJORS if m not in involved]
        
        quality_level = data.get("quality_level") or self.determine_quality_level(data)
        
        applicable_codes = data.get("applicable_codes") or []
        if design_stage == "初步设计":
            applicable_codes.append({"name": "水利水电工程初步设计报告编制规程", "code": "SL/T 619-2021"})
        elif design_stage == "可行性研究":
            applicable_codes.append({"name": "水利水电工程可行性研究报告编制规程", "code": "SL/T 618-2021"})
        
        seen_codes = set()
        unique_codes = []
        for code in applicable_codes:
            if code["code"] not in seen_codes:
                seen_codes.add(code["code"])
                unique_codes.append(code)
        
        review_method = "会议评审"
        verification_method = "产品校审"
        confirmation_method = "项目审查" if design_stage in ["初步设计", "实施方案"] else "项目咨询"
        
        return ISOProjectInfo(
            project_name=project_name,
            project_code=project_code,
            feature_code=feature_code,
            design_stage=design_stage,
            report_date=data.get("report_date"),
            client=data.get("client"),
            department="水利设计院",
            work_scope=data.get("work_scope"),
            engineering_overview=data.get("engineering_overview"),
            design_basis=data.get("design_basis"),
            technical_points=data.get("technical_points"),
            risk_points=data.get("risk_points"),
            customer_requirements=data.get("customer_requirements"),
            external_resources=data.get("external_resources"),
            quality_level=quality_level,
            project_grade=data.get("project_grade"),
            building_level=data.get("building_level"),
            flood_standard=data.get("flood_standard"),
            drainage_standard=data.get("drainage_standard"),
            involved_majors=involved,
            excluded_majors=excluded,
            applicable_codes=unique_codes[:10],
            design_review_method=review_method,
            design_verification_method=verification_method,
            design_confirmation_method=confirmation_method,
        )
    
    def _set_cell_text(self, cell, text: str):
        """设置单元格文本"""
        if not cell.paragraphs:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            run = para.add_run(text)
            run.font.size = Pt(10.5)
    
    def _check_checkbox(self, cell, checked: bool):
        """勾选复选框"""
        check_mark = "☑" if checked else "□"
        for para in cell.paragraphs:
            for run in para.runs:
                if "□" in run.text or "☑" in run.text:
                    run.text = run.text.replace("□", check_mark).replace("☑", check_mark)
                    return
    
    def fill_ty01_project_task(self, doc: Document, info: ISOProjectInfo):
        """填写TY01 项目任务书"""
        if len(doc.tables) < 1:
            return
        table = doc.tables[0]
        try:
            self._set_cell_text(table.rows[0].cells[1], info.project_name)
            self._set_cell_text(table.rows[1].cells[1], info.project_code)
            self._set_cell_text(table.rows[1].cells[3], info.feature_code)
            if info.work_scope:
                self._set_cell_text(table.rows[2].cells[1], info.work_scope)
            self._set_cell_text(table.rows[3].cells[1], info.department)
            filing_date = info.report_date or datetime.now().strftime("%Y年%m月")
            self._set_cell_text(table.rows[4].cells[1], filing_date)
            if info.client:
                self._set_cell_text(table.rows[5].cells[1], info.client)
            for cell in table.rows[7].cells:
                text = cell.text
                if info.quality_level == "A级" and "A级" in text:
                    self._check_checkbox(cell, True)
                elif info.quality_level == "B级" and "B级" in text:
                    self._check_checkbox(cell, True)
                elif info.quality_level == "C级" and "C级" in text:
                    self._check_checkbox(cell, True)
            risk_text = ""
            if info.customer_requirements:
                risk_text += f"顾客要求：{info.customer_requirements}\n"
            if info.risk_points:
                risk_text += f"主要风险要点：{info.risk_points}"
            if risk_text and len(table.rows) > 8:
                self._set_cell_text(table.rows[8].cells[1], risk_text)
        except Exception as e:
            print(f"填写TY01失败: {e}")
    
    def fill_ty02_project_plan(self, doc: Document, info: ISOProjectInfo):
        """填写TY02-1 工程项目策划表"""
        if len(doc.tables) < 2:
            return
        table = doc.tables[1]
        try:
            self._set_cell_text(table.rows[0].cells[1], info.project_name)
            if len(table.rows[0].cells) > 3:
                self._set_cell_text(table.rows[0].cells[3], info.project_code)
                self._set_cell_text(table.rows[0].cells[5], info.design_stage)
            if info.work_scope and len(table.rows) > 3:
                self._set_cell_text(table.rows[3].cells[1], info.work_scope)
            if info.engineering_overview and len(table.rows) > 4:
                self._set_cell_text(table.rows[4].cells[1], info.engineering_overview)
            if info.design_basis and len(table.rows) > 5:
                self._set_cell_text(table.rows[5].cells[1], info.design_basis)
            if info.technical_points and len(table.rows) > 6:
                self._set_cell_text(table.rows[6].cells[1], info.technical_points)
            if info.risk_points and len(table.rows) > 7:
                self._set_cell_text(table.rows[7].cells[1], info.risk_points)
            major_start_row = 20
            for i, major in enumerate(ALL_MAJORS):
                if major_start_row + i >= len(table.rows):
                    break
                row = table.rows[major_start_row + i]
                if major in info.excluded_majors:
                    for j in range(2, min(len(row.cells), 6)):
                        self._set_cell_text(row.cells[j], "/")
        except Exception as e:
            print(f"填写TY02-1失败: {e}")
    
    def fill_ty02_review_config(self, doc: Document, info: ISOProjectInfo):
        """填写TY02-2 项目校审配置表"""
        if len(doc.tables) < 3:
            return
        table = doc.tables[2]
        try:
            self._set_cell_text(table.rows[0].cells[1], info.project_name)
            if len(table.rows[0].cells) > 3:
                self._set_cell_text(table.rows[0].cells[3], info.design_stage)
                self._set_cell_text(table.rows[0].cells[5], info.quality_level)
            major_start_row = 4
            for i, major in enumerate(ALL_MAJORS):
                if major_start_row + i >= len(table.rows):
                    break
                row = table.rows[major_start_row + i]
                if major in info.excluded_majors:
                    for j in range(2, min(len(row.cells), 5)):
                        self._set_cell_text(row.cells[j], "/")
        except Exception as e:
            print(f"填写TY02-2失败: {e}")
    
    def fill_ty03_material_exchange(self, doc: Document, info: ISOProjectInfo):
        """填写TY03 互提资料单"""
        if len(doc.tables) < 4:
            return
        table = doc.tables[3]
        try:
            self._set_cell_text(table.rows[0].cells[1], info.project_name)
            if len(table.rows[0].cells) > 3:
                self._set_cell_text(table.rows[0].cells[3], info.design_stage)
        except Exception as e:
            print(f"填写TY03失败: {e}")
    
    def fill_ty04_product_card_major(self, doc: Document, info: ISOProjectInfo):
        """填写TY04-1 产品运行卡-专业"""
        if len(doc.tables) < 5:
            return
        table = doc.tables[4]
        try:
            product_name = f"{info.project_name} {info.design_stage}报告"
            self._set_cell_text(table.rows[0].cells[1], product_name)
            code_start_row = 3
            codes_to_fill = info.applicable_codes[:6]
            for i, code_item in enumerate(codes_to_fill):
                if code_start_row + i >= len(table.rows):
                    break
                code_text = f"《{code_item['name']}》{code_item['code']}"
                self._set_cell_text(table.rows[code_start_row + i].cells[1], code_text)
        except Exception as e:
            print(f"填写TY04-1失败: {e}")
    
    def fill_ty04_product_card_project(self, doc: Document, info: ISOProjectInfo):
        """填写TY04-2 产品运行卡-项目"""
        if len(doc.tables) < 6:
            return
        table = doc.tables[5]
        try:
            self._set_cell_text(table.rows[0].cells[1], info.project_name)
        except Exception as e:
            print(f"填写TY04-2失败: {e}")
    
    def fill_template(self, project_info: ISOProjectInfo, output_filename: Optional[str] = None) -> str:
        """填充Word模板"""
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"模板文件不存在: {self.template_path}")
        
        if not output_filename:
            output_filename = f"{project_info.project_code}_{project_info.project_name}_管理体系附表-设计部分.docx"
        output_filename = re.sub(r'[<>:"/\\|?*]', '_', output_filename)
        output_path = os.path.join(self.output_dir, output_filename)
        
        shutil.copy2(self.template_path, output_path)
        doc = Document(output_path)
        
        self.fill_ty01_project_task(doc, project_info)
        self.fill_ty02_project_plan(doc, project_info)
        self.fill_ty02_review_config(doc, project_info)
        self.fill_ty03_material_exchange(doc, project_info)
        self.fill_ty04_product_card_major(doc, project_info)
        self.fill_ty04_product_card_project(doc, project_info)
        
        doc.save(output_path)
        return output_path
    
    async def generate_document(self, file_path: str, filename: str,
                               supplementary_info: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """生成ISO文档"""
        task_id = str(uuid.uuid4())
        
        from app.services.document_parser import document_parser
        text_content = ""
        if filename.endswith('.docx') or filename.endswith('.doc'):
            text_content = document_parser._parse_docx(file_path)
        elif filename.endswith('.pdf'):
            text_content = document_parser._parse_pdf(file_path)
        else:
            text_content = document_parser._parse_txt(file_path)
        
        project_info = await self.extract_project_info(text_content, filename)
        
        if supplementary_info:
            for key, value in supplementary_info.items():
                if hasattr(project_info, key) and value:
                    setattr(project_info, key, value)
        
        warnings = []
        if not project_info.client:
            warnings.append("未能自动识别业主单位，请人工补充")
        if not project_info.report_date:
            warnings.append("未能自动识别报告日期，请人工确认")
        warnings.append(f"质量控制分级自动判定为{project_info.quality_level}，请人工复核")
        
        output_path = self.fill_template(project_info)
        
        self.tasks[task_id] = {
            "task_id": task_id,
            "status": "completed",
            "project_info": project_info,
            "output_file": output_path,
            "created_at": datetime.now().isoformat(),
        }
        
        download_url = f"/api/iso/download/{task_id}"
        
        return {
            "task_id": task_id,
            "status": "completed",
            "project_info": project_info.model_dump(),
            "output_file": output_path,
            "download_url": download_url,
            "warnings": warnings,
            "message": "文档生成成功，请确认信息后下载",
        }
    
    def get_task(self, task_id: str) -> Optional[Dict[str, Any]]:
        return self.tasks.get(task_id)
    
    def get_output_file(self, task_id: str) -> Optional[str]:
        task = self.tasks.get(task_id)
        if task and os.path.exists(task.get("output_file", "")):
            return task["output_file"]
        return None
    
    def regenerate_with_confirmation(self, task_id: str, project_info_dict: Dict[str, Any]) -> Optional[str]:
        try:
            project_info = ISOProjectInfo(**project_info_dict)
            output_path = self.fill_template(project_info)
            if task_id in self.tasks:
                self.tasks[task_id]["project_info"] = project_info
                self.tasks[task_id]["output_file"] = output_path
            return output_path
        except Exception as e:
            print(f"重新生成失败: {e}")
            return None


iso_service = ISOService()
