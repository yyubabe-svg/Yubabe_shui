"""
报告章节智能生成服务
- 基于章节模板 + 参考文档 + 参数，生成报告章节
- 支持大纲生成、SSE 流式逐段生成
- 支持段落 accept/edit/regenerate
- 支持导出 Word（.docx）
"""
import os
import re
import json
import asyncio
from typing import Optional, List, Dict, Any, AsyncGenerator
from datetime import datetime
from sqlalchemy.orm import Session

from app.models.report_section import (
    ReportSectionTemplate, ReportSectionTask, ReportSectionDraft,
    SectionTaskStatus, DraftStatus, ParagraphType,
)
from app.models.document import Document, Chunk
from app.schemas.report_section import (
    OutlineItem, OutlineResult, GeneratedParagraph,
)
from app.services.llm_service import llm_service
from app.services.vector_store import vector_store
from app.services.embedding import embedding_service
from app.core.config import settings


# ========================================================================
# Prompt 模板
# ========================================================================

OUTLINE_SYSTEM_PROMPT = """你是"蜀水智库 AI"的水利工程设计报告撰写助手。
请根据给定的章节模板要求、项目参数和参考资料，生成该章节的详细写作大纲。
输出必须是严格的 JSON，不要附加任何解释文字。"""

OUTLINE_USER_PROMPT = """请为以下章节生成详细写作大纲。

【章节信息】
章节编号：{chapter_number}
章节标题：{chapter_title}
写作指导：{writing_prompt}
适用项目类型：{applicable_types}
适用设计阶段：{applicable_stages}

【项目参数】
{params_text}

【必需参数】
{required_params}

【参考资料关键词】
{reference_keywords}

请严格按以下 JSON 结构输出：
{{
  "chapter_title": "{chapter_title}",
  "chapter_number": "{chapter_number}",
  "items": [
    {{
      "paragraph_id": "heading_1",
      "parent_paragraph_id": null,
      "paragraph_type": "heading",
      "level": 2,
      "title": "小节标题",
      "writing_instruction": "该段落应包含的核心要点",
      "keywords": ["关键词1", "关键词2"]
    }},
    {{
      "paragraph_id": "para_1",
      "parent_paragraph_id": "heading_1",
      "paragraph_type": "paragraph",
      "level": null,
      "title": "段落概要",
      "writing_instruction": "应说明的内容要点",
      "keywords": ["关键词"]
    }}
  ]
}}

注意：
1. paragraph_id 唯一且按顺序编号（heading_N / para_N / list_N / table_N）；
2. paragraph_type 只能是 heading/paragraph/list/table 之一；
3. heading 类型需要填写 level（2/3/4）；
4. 大纲应覆盖章节模板要求的全部内容要点。
"""

PARAGRAPH_SYSTEM_PROMPT = """你是"蜀水智库 AI"的水利工程设计报告撰写助手。
请严格依据提供的参考资料和写作要求，撰写专业的报告段落。
要求：
1. 使用专业、严谨的水利工程设计报告语言风格；
2. 所有技术参数和结论必须有参考资料依据，不得编造；
3. 如果参考资料不足，请在文中注明"需补充"而不是编造数据；
4. 使用标准的工程术语和规范表述；
5. 直接输出段落正文，不要输出标题、解释或markdown标记。"""

PARAGRAPH_USER_PROMPT = """请撰写报告章节中的一个段落。

【章节】{chapter_number} {chapter_title}
【段落类型】{paragraph_type}
【段落概要】{title}
【写作要求】{writing_instruction}
【父节标题】{parent_title}

【项目参数】
{params_text}

【已撰写的上下文内容】
{context_text}

【参考资料】
{references}

请直接输出该段落的正文内容（不要输出标题和说明）：
"""

REGENERATE_USER_PROMPT = """请对以下段落进行修改重写。

【原段落内容】
{original_content}

【用户反馈/修改要求】
{feedback}

【章节上下文】
{context_text}

【参考资料】
{references}

请直接输出修改后的段落正文：
"""


# ========================================================================
# SectionGenerator
# ========================================================================

class SectionGenerator:
    """报告章节智能生成器"""

    def __init__(self, db: Session):
        self.db = db

    # ------------------------------------------------------------------
    # 任务管理
    # ------------------------------------------------------------------

    def create_task(
        self,
        project_id: int,
        template_id: int,
        doc_ids: Optional[List[int]] = None,
        params: Optional[Dict[str, Any]] = None,
        created_by: Optional[str] = None,
    ) -> ReportSectionTask:
        """创建章节生成任务"""
        tpl = self.db.query(ReportSectionTemplate).filter(
            ReportSectionTemplate.id == template_id
        ).first()
        if not tpl:
            raise ValueError(f"章节模板 {template_id} 不存在")

        # 合并模板必需参数和用户参数
        merged_params = {}
        if tpl.required_params:
            for p in tpl.required_params:
                if isinstance(p, dict):
                    merged_params[p.get("name", "")] = p.get("default", "")
                elif isinstance(p, str):
                    merged_params[p] = ""
        if params:
            merged_params.update(params)

        task = ReportSectionTask(
            project_id=project_id,
            template_id=template_id,
            document_ids=doc_ids or [],
            status=SectionTaskStatus.PENDING.value,
            params_override=merged_params,
            progress=0,
            created_by=created_by,
        )
        self.db.add(task)
        self.db.commit()
        self.db.refresh(task)
        return task

    def get_task(self, task_id: int) -> Optional[ReportSectionTask]:
        return self.db.query(ReportSectionTask).filter(ReportSectionTask.id == task_id).first()

    def get_task_with_details(self, task_id: int) -> Optional[ReportSectionTask]:
        from sqlalchemy.orm import joinedload
        return (
            self.db.query(ReportSectionTask)
            .options(
                joinedload(ReportSectionTask.drafts),
                joinedload(ReportSectionTask.template),
            )
            .filter(ReportSectionTask.id == task_id)
            .first()
        )

    # ------------------------------------------------------------------
    # 大纲生成
    # ------------------------------------------------------------------

    async def generate_outline(self, task_id: int) -> OutlineResult:
        """生成章节大纲"""
        task = self.get_task_with_details(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")

        tpl = task.template
        params = task.params_override or {}
        params_text = "\n".join(f"- {k}: {v}" for k, v in params.items()) if params else "（暂无）"
        required_params = (
            "\n".join(f"- {p}" for p in tpl.required_params)
            if tpl.required_params else "（无特殊要求）"
        )
        keywords = (
            "、".join(tpl.reference_keywords)
            if tpl.reference_keywords else "（无）"
        )

        prompt = OUTLINE_USER_PROMPT.format(
            chapter_number=tpl.chapter_number,
            chapter_title=tpl.title,
            writing_prompt=tpl.writing_prompt or "按照水利水电工程初步设计报告编制规程撰写",
            applicable_types=", ".join(tpl.applicable_project_types or []) or "通用",
            applicable_stages=", ".join(tpl.applicable_stages or []) or "通用",
            params_text=params_text,
            required_params=required_params,
            reference_keywords=keywords,
        )

        messages = [
            {"role": "system", "content": OUTLINE_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ]

        raw = await llm_service.chat(messages, temperature=0.3, max_tokens=4096)
        outline = self._parse_outline_json(raw, tpl.chapter_number, tpl.title)

        # 保存大纲
        task.status = SectionTaskStatus.OUTLINE_GENERATING.value
        task.outline_json = outline.model_dump()
        task.status = SectionTaskStatus.GENERATING.value
        self.db.commit()

        # 创建空的草稿段落占位
        self._create_draft_placeholders(task, outline)

        return outline

    def update_outline(self, task_id: int, outline_json: Dict[str, Any]) -> ReportSectionTask:
        """手动编辑大纲"""
        task = self.get_task(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")
        task.outline_json = outline_json
        # 重新创建占位段落
        self._clear_drafts(task_id)
        outline = OutlineResult(**outline_json)
        self._create_draft_placeholders(task, outline)
        self.db.commit()
        self.db.refresh(task)
        return task

    def _create_draft_placeholders(self, task: ReportSectionTask, outline: OutlineResult):
        """根据大纲创建空的草稿段落记录"""
        # 先添加章节标题作为 heading
        main_heading = ReportSectionDraft(
            task_id=task.id,
            paragraph_id="heading_0",
            parent_paragraph_id=None,
            paragraph_type=ParagraphType.HEADING.value,
            level=int(task.template.level) if task.template and task.template.level else 1,
            content=f"{task.template.chapter_number} {task.template.title}" if task.template else outline.chapter_title,
            status=DraftStatus.ACCEPTED.value,
            sort_order=0,
        )
        self.db.add(main_heading)

        for idx, item in enumerate(outline.items, 1):
            draft = ReportSectionDraft(
                task_id=task.id,
                paragraph_id=item.paragraph_id,
                parent_paragraph_id=item.parent_paragraph_id,
                paragraph_type=item.paragraph_type,
                level=item.level,
                content="",  # 待生成
                status=DraftStatus.PENDING.value,
                sources_json=None,
                sort_order=idx,
                feedback=None,
            )
            self.db.add(draft)
        self.db.flush()

    def _clear_drafts(self, task_id: int):
        """清除任务的所有草稿"""
        self.db.query(ReportSectionDraft).filter(
            ReportSectionDraft.task_id == task_id
        ).delete()
        self.db.flush()

    def _parse_outline_json(self, raw: str, chapter_number: str, chapter_title: str) -> OutlineResult:
        """解析大纲 JSON"""
        cleaned = raw.strip()
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r'\{[\s\S]*\}', cleaned)
            if not match:
                return OutlineResult(chapter_title=chapter_title, chapter_number=chapter_number, items=[])
            try:
                data = json.loads(match.group(0))
            except json.JSONDecodeError:
                return OutlineResult(chapter_title=chapter_title, chapter_number=chapter_number, items=[])

        items = []
        for item_data in data.get("items", []):
            try:
                items.append(OutlineItem(**item_data))
            except Exception:
                continue

        return OutlineResult(
            chapter_title=data.get("chapter_title", chapter_title),
            chapter_number=data.get("chapter_number", chapter_number),
            items=items,
        )

    # ------------------------------------------------------------------
    # SSE 流式逐段生成
    # ------------------------------------------------------------------

    async def generate_section(self, task_id: int, start_from: Optional[str] = None) -> AsyncGenerator[Dict[str, Any], None]:
        """
        SSE 流式逐段生成，逐段落 yield 事件。
        事件: outline / paragraph_start / paragraph_delta / paragraph_done / done / error
        """
        task = self.get_task_with_details(task_id)
        if not task:
            yield {"event": "error", "error": f"任务 {task_id} 不存在"}
            return

        try:
            # 1. 如果还没有大纲，先生成
            if not task.outline_json:
                yield {"event": "outline", "task_id": task_id, "message": "正在生成大纲..."}
                outline = await self.generate_outline(task_id)
                yield {"event": "outline", "task_id": task_id, "outline": outline.model_dump()}
                self.db.refresh(task)

            outline = OutlineResult(**task.outline_json)
            drafts = {d.paragraph_id: d for d in task.drafts}

            # 2. 检索参考文档内容
            references = await self._retrieve_references(task, outline)

            # 3. 确定需要生成的段落
            items_to_generate = []
            started = start_from is None
            for item in outline.items:
                if not started and item.paragraph_id == start_from:
                    started = True
                if started:
                    draft = drafts.get(item.paragraph_id)
                    if draft and draft.status in (DraftStatus.PENDING.value, DraftStatus.REGENERATED.value):
                        items_to_generate.append(item)

            total = len(items_to_generate)
            if total == 0:
                yield {
                    "event": "done", "task_id": task_id,
                    "message": "所有段落已生成", "progress": 100,
                }
                return

            # 4. 逐段流式生成
            context_parts: List[str] = []
            for idx, item in enumerate(items_to_generate):
                # 收集已接受/已生成的上下文
                context_text = self._build_context(task, item.paragraph_id)

                yield {
                    "event": "paragraph_start",
                    "task_id": task_id,
                    "paragraph_id": item.paragraph_id,
                    "paragraph_type": item.paragraph_type,
                    "title": item.title,
                    "progress": int(idx / total * 100),
                }

                # 组装 prompt
                params = task.params_override or {}
                params_text = "\n".join(f"- {k}: {v}" for k, v in params.items()) if params else "（暂无）"

                prompt = PARAGRAPH_USER_PROMPT.format(
                    chapter_number=outline.chapter_number,
                    chapter_title=outline.chapter_title,
                    paragraph_type=item.paragraph_type,
                    title=item.title,
                    writing_instruction=item.writing_instruction or "",
                    parent_title=self._get_parent_title(outline, item),
                    params_text=params_text,
                    context_text=context_text[-2000:] if context_text else "（暂无）",
                    references=references[:3000] if references else "（暂无参考资料）",
                )

                messages = [
                    {"role": "system", "content": PARAGRAPH_SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ]

                # 流式生成
                content_parts = []
                async for chunk in llm_service.chat_stream(messages, temperature=0.3):
                    content_parts.append(chunk)
                    yield {
                        "event": "paragraph_delta",
                        "task_id": task_id,
                        "paragraph_id": item.paragraph_id,
                        "delta": chunk,
                    }

                full_content = "".join(content_parts).strip()

                # 保存段落
                draft = drafts.get(item.paragraph_id)
                if draft:
                    draft.content = full_content
                    draft.status = DraftStatus.GENERATED.value
                    draft.sources_json = []  # 可扩展为记录引用来源
                    self.db.flush()

                yield {
                    "event": "paragraph_done",
                    "task_id": task_id,
                    "paragraph_id": item.paragraph_id,
                    "content": full_content,
                    "progress": int((idx + 1) / total * 100),
                }

            # 5. 汇总
            self._assemble_content(task)
            task.status = SectionTaskStatus.EDITING.value
            task.progress = 100
            self.db.commit()

            yield {
                "event": "done",
                "task_id": task_id,
                "progress": 100,
                "message": f"章节生成完成，共 {total} 个段落",
            }

        except Exception as e:
            task.status = SectionTaskStatus.FAILED.value
            task.error_message = str(e)
            self.db.commit()
            yield {"event": "error", "task_id": task_id, "error": str(e)}

    # ------------------------------------------------------------------
    # 段落操作：accept / edit / regenerate
    # ------------------------------------------------------------------

    def accept_paragraph(self, task_id: int, paragraph_id: str) -> Optional[ReportSectionDraft]:
        """接受段落"""
        draft = self._get_draft(task_id, paragraph_id)
        if not draft:
            return None
        draft.status = DraftStatus.ACCEPTED.value
        self.db.commit()
        self.db.refresh(draft)
        return draft

    def edit_paragraph(self, task_id: int, paragraph_id: str, content: str, note: Optional[str] = None) -> Optional[ReportSectionDraft]:
        """编辑段落内容"""
        draft = self._get_draft(task_id, paragraph_id)
        if not draft:
            return None
        draft.content = content
        draft.status = DraftStatus.EDITED.value
        if note:
            draft.feedback = note
        self._assemble_content_for_task(task_id)
        self.db.commit()
        self.db.refresh(draft)
        return draft

    async def regenerate_paragraph(self, task_id: int, paragraph_id: str, feedback: Optional[str] = None) -> AsyncGenerator[Dict[str, Any], None]:
        """重新生成段落（流式）"""
        task = self.get_task_with_details(task_id)
        if not task:
            yield {"event": "error", "error": f"任务 {task_id} 不存在"}
            return

        draft = self._get_draft(task_id, paragraph_id)
        if not draft:
            yield {"event": "error", "error": f"段落 {paragraph_id} 不存在"}
            return

        if feedback:
            draft.feedback = feedback

        outline = OutlineResult(**task.outline_json) if task.outline_json else None
        if not outline:
            yield {"event": "error", "error": "大纲不存在，请先生成大纲"}
            return

        item = next((i for i in outline.items if i.paragraph_id == paragraph_id), None)
        if not item:
            yield {"event": "error", "error": f"大纲中未找到段落 {paragraph_id}"}
            return

        references = await self._retrieve_references(task, outline)
        context_text = self._build_context(task, paragraph_id)
        params = task.params_override or {}
        params_text = "\n".join(f"- {k}: {v}" for k, v in params.items()) if params else "（暂无）"

        if feedback and draft.content:
            prompt = REGENERATE_USER_PROMPT.format(
                original_content=draft.content,
                feedback=feedback,
                context_text=context_text[-2000:],
                references=references[:3000] if references else "（暂无）",
            )
        else:
            prompt = PARAGRAPH_USER_PROMPT.format(
                chapter_number=outline.chapter_number,
                chapter_title=outline.chapter_title,
                paragraph_type=item.paragraph_type,
                title=item.title,
                writing_instruction=item.writing_instruction or "",
                parent_title=self._get_parent_title(outline, item),
                params_text=params_text,
                context_text=context_text[-2000:] if context_text else "（暂无）",
                references=references[:3000] if references else "（暂无）",
            )

        messages = [
            {"role": "system", "content": PARAGRAPH_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ]

        yield {"event": "paragraph_start", "task_id": task_id, "paragraph_id": paragraph_id}

        content_parts = []
        async for chunk in llm_service.chat_stream(messages, temperature=0.4):
            content_parts.append(chunk)
            yield {"event": "paragraph_delta", "task_id": task_id, "paragraph_id": paragraph_id, "delta": chunk}

        full_content = "".join(content_parts).strip()
        draft.content = full_content
        draft.status = DraftStatus.REGENERATED.value
        self.db.flush()
        self._assemble_content(task)
        self.db.commit()

        yield {
            "event": "paragraph_done",
            "task_id": task_id,
            "paragraph_id": paragraph_id,
            "content": full_content,
        }

    def _get_draft(self, task_id: int, paragraph_id: str) -> Optional[ReportSectionDraft]:
        return self.db.query(ReportSectionDraft).filter(
            ReportSectionDraft.task_id == task_id,
            ReportSectionDraft.paragraph_id == paragraph_id,
        ).first()

    def _get_parent_title(self, outline: OutlineResult, item: OutlineItem) -> str:
        if not item.parent_paragraph_id:
            return ""
        parent = next((i for i in outline.items if i.paragraph_id == item.parent_paragraph_id), None)
        return parent.title if parent else ""

    def _build_context(self, task: ReportSectionTask, current_paragraph_id: str) -> str:
        """构建已生成段落的上下文"""
        parts = []
        for d in sorted(task.drafts, key=lambda x: x.sort_order):
            if d.paragraph_id == current_paragraph_id:
                break
            if d.status in (DraftStatus.GENERATED.value, DraftStatus.ACCEPTED.value, DraftStatus.EDITED.value, DraftStatus.REGENERATED.value) and d.content:
                if d.paragraph_type == ParagraphType.HEADING.value:
                    prefix = "#" * (d.level or 2) + " "
                    parts.append(f"{prefix}{d.content}")
                elif d.paragraph_type == ParagraphType.LIST.value:
                    parts.append(d.content)
                else:
                    parts.append(d.content)
        return "\n\n".join(parts)

    # ------------------------------------------------------------------
    # 参考资料检索
    # ------------------------------------------------------------------

    async def _retrieve_references(self, task: ReportSectionTask, outline: OutlineResult) -> str:
        """从参考文档和规范库检索相关内容"""
        doc_ids = task.document_ids or []
        all_refs = []

        # 收集关键词
        all_keywords = []
        if task.template and task.template.reference_keywords:
            all_keywords.extend(task.template.reference_keywords)
        for item in outline.items:
            if item.keywords:
                all_keywords.extend(item.keywords)
        if task.template and task.template.title:
            all_keywords.append(task.template.title)

        query_text = " ".join(all_keywords[:10])
        if not query_text:
            return ""

        try:
            embedding = await asyncio.to_thread(embedding_service.embed_query, query_text)
            results = vector_store.search(embedding, top_k=8)

            # 如果指定了文档ID，优先使用这些文档
            if doc_ids:
                doc_chunks = (
                    self.db.query(Chunk)
                    .filter(Chunk.document_id.in_(doc_ids))
                    .limit(20)
                    .all()
                )
                for chunk in doc_chunks:
                    all_refs.append({
                        "file_name": chunk.document.title if hasattr(chunk, 'document') and chunk.document else "参考文档",
                        "page_number": getattr(chunk, 'page_number', ''),
                        "text": (chunk.chunk_text or '')[:500],
                    })

            for r in results:
                meta = r.get("metadata", {})
                all_refs.append({
                    "file_name": meta.get("file_name", "未知"),
                    "page_number": meta.get("page_number", ""),
                    "text": meta.get("text", "")[:500],
                })
        except Exception as e:
            print(f"[SectionGenerator] 参考资料检索失败: {e}")

        parts = []
        for i, ref in enumerate(all_refs[:10], 1):
            part = f"[{i}] 文件：{ref['file_name']}\n"
            if ref.get("page_number"):
                part += f"页码：{ref['page_number']}\n"
            part += f"内容：{ref['text']}\n"
            parts.append(part)
        return "\n---\n".join(parts)

    # ------------------------------------------------------------------
    # 内容汇总
    # ------------------------------------------------------------------

    def _assemble_content(self, task: ReportSectionTask):
        """将所有已接受/已生成段落汇总为完整内容"""
        parts = []
        for d in sorted(task.drafts, key=lambda x: x.sort_order):
            if d.content and d.status not in (DraftStatus.PENDING.value,):
                if d.paragraph_type == ParagraphType.HEADING.value:
                    level = d.level or 2
                    parts.append(f"\n{'#' * level} {d.content}\n")
                elif d.paragraph_type == ParagraphType.LIST.value:
                    parts.append(d.content)
                elif d.paragraph_type == ParagraphType.TABLE.value:
                    parts.append(d.content)
                else:
                    parts.append(d.content + "\n")
        task.assembled_content = "\n".join(parts)
        task.progress = self._calc_progress(task)
        self.db.flush()

    def _assemble_content_for_task(self, task_id: int):
        task = self.get_task_with_details(task_id)
        if task:
            self._assemble_content(task)
            self.db.flush()

    def _calc_progress(self, task: ReportSectionTask) -> int:
        drafts = task.drafts
        if not drafts:
            return 0
        done = sum(1 for d in drafts if d.status in (
            DraftStatus.GENERATED.value, DraftStatus.ACCEPTED.value,
            DraftStatus.EDITED.value, DraftStatus.REGENERATED.value,
        ))
        return int(done / len(drafts) * 100)

    # ------------------------------------------------------------------
    # 导出 Word
    # ------------------------------------------------------------------

    def export_docx(self, task_id: int) -> Dict[str, Any]:
        """导出章节为 Word 文件"""
        task = self.get_task_with_details(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")

        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx 未安装，请执行 pip install python-docx")

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)
        style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
        style.paragraph_format.line_spacing = 1.5

        # 按 sort_order 输出段落
        for d in sorted(task.drafts, key=lambda x: x.sort_order):
            if not d.content:
                continue

            if d.paragraph_type == ParagraphType.HEADING.value:
                level = d.level or 1
                level = min(level, 4)  # docx 支持 heading 1-4
                heading = doc.add_heading(d.content, level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in heading.runs:
                    run.font.name = "黑体"
            elif d.paragraph_type == ParagraphType.LIST.value:
                for line in d.content.split("\n"):
                    line = line.strip()
                    if line:
                        p = doc.add_paragraph(line, style="List Bullet")
                        p.paragraph_format.first_line_indent = Cm(0)
            elif d.paragraph_type == ParagraphType.TABLE.value:
                # 简单支持 markdown 表格，否则直接输出
                self._add_table_or_text(doc, d.content)
            else:
                p = doc.add_paragraph(d.content)

        # 保存
        export_dir = os.path.join(settings.UPLOAD_DIR or "./uploads", "section_drafts")
        os.makedirs(export_dir, exist_ok=True)
        tpl = task.template
        chap_label = f"{tpl.chapter_number}_{tpl.title}" if tpl else f"章节{task_id}"
        filename = f"{chap_label}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        task.output_file_path = filepath
        task.output_filename = filename
        task.status = SectionTaskStatus.COMPLETED.value
        task.completed_at = datetime.utcnow()
        self.db.commit()

        para_count = sum(
            1 for d in task.drafts
            if d.content and d.status != DraftStatus.PENDING.value
        )
        return {
            "task_id": task_id,
            "file_path": filepath,
            "file_name": filename,
            "paragraph_count": para_count,
            "exported_at": datetime.utcnow(),
        }

    def _add_table_or_text(self, doc: "DocxDocument", content: str):
        """尝试解析 markdown 表格，否则按普通文本输出"""
        lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
        if len(lines) >= 2 and "|" in lines[0] and re.match(r'^[\s|:-]+$', lines[1]):
            # 解析 markdown 表格
            headers = [c.strip() for c in lines[0].strip("|").split("|")]
            rows = []
            for line in lines[2:]:
                cells = [c.strip() for c in line.strip("|").split("|")]
                rows.append(cells)
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for r, row in enumerate(rows, 1):
                for c, val in enumerate(row):
                    if c < len(headers):
                        table.rows[r].cells[c].text = val
        else:
            doc.add_paragraph(content)


# ========================================================================
# 工厂函数
# ========================================================================

def get_section_generator(db: Session) -> SectionGenerator:
    return SectionGenerator(db)
