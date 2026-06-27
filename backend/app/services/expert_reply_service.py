"""
专家意见回复服务
- 解析审查意见文档（正则 + LLM 切分意见）
- RAG 检索规范+报告原文，生成回复
- 区分 已修改 / 已补充 / 解释说明
- 导出 Word 回复表（逐条意见回复）
"""
import os
import re
import json
import asyncio
from typing import Optional, List, Dict, Any, AsyncGenerator
from datetime import datetime
from sqlalchemy.orm import Session

from app.models.expert_reply import (
    ExpertReplyTask, ExpertOpinion, ExpertReplyItem,
    ReplyTaskStatus, OpinionType, MajorCategory, ModifyStatus, ReplyStatus,
)
from app.models.document import Document, Chunk
from app.schemas.expert_reply import (
    ParsedOpinionItem, ReplyGenerateResult,
    MAJOR_CATEGORY_MAP, MODIFY_STATUS_OPTIONS,
)
from app.services.llm_service import llm_service
from app.services.vector_store import vector_store
from app.services.embedding import embedding_service
from app.core.config import settings


# ========================================================================
# 正则切分模式
# ========================================================================

# 常见意见编号模式：
#   1. / 1、 / 1. / （1） / 【1】 / 一、 / 1)
OPINION_INDEX_PATTERNS = [
    re.compile(r'^\s*(\d+)[、．.\s]+', re.MULTILINE),
    re.compile(r'^\s*[（(](\d+)[）)]\s*', re.MULTILINE),
    re.compile(r'^\s*[【\[](\d+)[】\]]\s*', re.MULTILINE),
    re.compile(r'^\s*(\d+)\)\s*', re.MULTILINE),
]

# 专家姓名模式（常见格式："张三（水工）"、"李四:")
EXPERT_NAME_PATTERN = re.compile(
    r'(?:专家|审查人|发言人)?[：:]?\s*([\u4e00-\u9fa5]{2,4})\s*'
    r'(?:[（(]([\u4e00-\u9fa5]{2,6})[）)])?\s*[：:]',
)

# 专业分类关键词
CATEGORY_KEYWORDS = {
    "hydrology": ["水文", "洪水", "径流", "暴雨", "频率", "洪峰", "水位", "流量"],
    "geology": ["地质", "勘察", "岩性", "断层", "岩溶", "渗漏", "稳定", "承载力", "土壤"],
    "hydraulic": ["水工", "坝", "堤", "溢洪道", "闸", "隧洞", "消能", "防渗", "护坡", "结构", "稳定计算"],
    "construction": ["施工", "导流", "围堰", "度汛", "工期", "浇筑", "开挖", "填筑"],
    "investment": ["投资", "概算", "预算", "造价", "工程量", "定额", "费用"],
    "soil_conservation": ["水保", "水土保持", "弃渣", "植被恢复", "水土流失"],
    "environment": ["环境", "生态", "水质", "环保", "影响评价", "鱼类", "洄游"],
    "electromechanical": ["机电", "金结", "金属结构", "闸门", "启闭机", "电气", "机组"],
    "resettlement": ["移民", "占地", "征地", "搬迁", "安置"],
    "format": ["格式", "文字", "编号", "编排", "错字", "标点", "排版"],
}


# ========================================================================
# Prompt 模板
# ========================================================================

PARSE_SYSTEM_PROMPT = """你是"蜀水智库 AI"的审查意见解析助手。
请将审查意见文本切分为独立的意见条目，并识别专家姓名、专业分类和意见类型。
输出必须是严格的 JSON 数组，不要附加任何解释文字。"""

PARSE_USER_PROMPT = """请解析以下审查意见文本，切分为独立的意见条目。

【意见全文】
{text}

【已有切分提示】
初步检测到可能的意见编号：{pre_detected_indices}

请输出 JSON 数组，每条意见格式如下：
[
  {{
    "opinion_index": 1,
    "expert_name": "专家姓名或null",
    "major_category": "专业分类key(hydrology/geology/hydraulic/construction/investment/soil_conservation/environment/electromechanical/resettlement/format/other)",
    "opinion_type": "modify/supplement/explain/other",
    "content": "意见完整内容",
    "page_number": 页码或null,
    "chapter_path": "涉及章节或null"
  }}
]

规则：
1. opinion_index 按顺序编号，从1开始；
2. 如果文本中无法识别专家姓名，expert_name 设为 null；
3. opinion_type 判断：要求修改已有内容=modify，要求增加内容=supplement，仅需解释说明=explain，其余=other；
4. major_category 根据内容关键词判断，无法判断时设为 other；
5. 不要遗漏任何一条意见。
"""

REPLY_SYSTEM_PROMPT = """你是"蜀水智库 AI"的水利工程设计报告专家意见回复助手。
请根据专家意见、报告原文和相关规范条文，撰写正式的回复内容。
要求：
1. 回复语气正式、专业、礼貌，符合设计院回复审查意见的行文规范；
2. 根据意见性质，明确回复类型：
   - 已修改：意见合理，已按意见修改报告内容；
   - 已补充：意见要求补充内容，已补充；
   - 解释说明：意见涉及问题已有合理处理或不需修改，给出解释；
   - 不修改：意见不采纳，说明理由（谨慎使用）；
3. 修改/补充类回复应说明具体修改位置（章节/页码）和修改内容；
4. 解释类回复应引用规范条文或设计依据；
5. 所有结论必须有依据，不得编造；
6. 输出必须是严格的 JSON，不要附加任何解释文字。"""

REPLY_USER_PROMPT = """请根据以下信息撰写专家意见回复。

【专家意见】
序号：{opinion_index}
专家：{expert_name}
专业：{major_category}
意见内容：{opinion_content}
意见涉及章节：{chapter_path}
意见涉及页码：{page_number}

【报告相关原文】
{report_context}

【检索到的规范依据】
{standards_context}

请严格按以下 JSON 格式输出：
{{
  "reply_content": "回复正文（150-500字，正式书面语）",
  "modify_status": "已修改/已补充/解释说明/不修改",
  "modify_location": "修改位置描述（如'第3章 3.2节 堤顶高程计算'），解释类可填null",
  "modify_page": "修改页码，如'P12-15'，解释类可填null",
  "sources": [
    {{"type": "规范/报告", "name": "名称", "page": "页码", "snippet": "引用片段"}}
  ]
}}
"""


# ========================================================================
# ExpertReplyService
# ========================================================================

class ExpertReplyService:
    """专家意见回复服务"""

    def __init__(self, db: Session):
        self.db = db

    # ------------------------------------------------------------------
    # 任务管理
    # ------------------------------------------------------------------

    def create_task(
        self,
        project_id: int,
        opinion_doc_id: int,
        report_doc_id: Optional[int] = None,
        meeting_name: Optional[str] = None,
        meeting_date: Optional[str] = None,
        created_by: Optional[str] = None,
    ) -> ExpertReplyTask:
        """创建专家意见回复任务"""
        opinion_doc = self.db.query(Document).filter(Document.id == opinion_doc_id).first()
        if not opinion_doc:
            raise ValueError(f"意见文档 {opinion_doc_id} 不存在")

        if report_doc_id:
            report_doc = self.db.query(Document).filter(Document.id == report_doc_id).first()
            if not report_doc:
                raise ValueError(f"报告文档 {report_doc_id} 不存在")

        task = ExpertReplyTask(
            project_id=project_id,
            opinion_document_id=opinion_doc_id,
            report_document_id=report_doc_id,
            status=ReplyTaskStatus.PENDING.value,
            meeting_name=meeting_name,
            meeting_date=meeting_date,
            progress=0,
            created_by=created_by,
        )
        self.db.add(task)
        self.db.commit()
        self.db.refresh(task)
        return task

    def get_task(self, task_id: int) -> Optional[ExpertReplyTask]:
        return self.db.query(ExpertReplyTask).filter(ExpertReplyTask.id == task_id).first()

    def get_task_with_details(self, task_id: int) -> Optional[ExpertReplyTask]:
        from sqlalchemy.orm import joinedload
        return (
            self.db.query(ExpertReplyTask)
            .options(
                joinedload(ExpertReplyTask.opinions).joinedload(ExpertOpinion.reply_item),
            )
            .filter(ExpertReplyTask.id == task_id)
            .first()
        )

    # ------------------------------------------------------------------
    # 意见解析
    # ------------------------------------------------------------------

    async def parse_opinions(self, task_id: int) -> AsyncGenerator[Dict[str, Any], None]:
        """
        解析意见文档，切分独立意见条目（异步生成器）。
        事件: progress / opinion_parsed / done / error
        """
        task = self.get_task(task_id)
        if not task:
            yield {"event": "error", "error": f"任务 {task_id} 不存在"}
            return

        task.status = ReplyTaskStatus.PARSING.value
        task.progress = 0
        self.db.commit()

        try:
            # 1. 加载意见文档全文
            text = await self._load_document_text(task.opinion_document_id)
            if not text:
                raise ValueError("意见文档内容为空")

            yield {"event": "progress", "task_id": task_id, "progress": 10, "message": "文档加载完成"}

            # 2. 正则预切分
            pre_indices = self._pre_detect_indices(text)

            yield {"event": "progress", "task_id": task_id, "progress": 20, "message": f"预检测到 {len(pre_indices)} 个编号"}

            # 3. LLM 结构化切分
            parsed_items = await self._llm_parse_opinions(text, pre_indices)

            yield {"event": "progress", "task_id": task_id, "progress": 50, "message": f"LLM解析完成，共 {len(parsed_items)} 条意见"}

            # 4. 保存意见
            total = len(parsed_items)
            expert_set = set()
            for idx, item in enumerate(parsed_items):
                opinion = ExpertOpinion(
                    reply_task_id=task.id,
                    opinion_index=item.opinion_index,
                    order_index=idx,
                    expert_name=item.expert_name,
                    major_category=item.major_category or "other",
                    opinion_type=item.opinion_type or "other",
                    content=item.content,
                    page_number=item.page_number,
                    chapter_path=item.chapter_path,
                )
                self.db.add(opinion)
                if item.expert_name:
                    expert_set.add(item.expert_name)

                yield {
                    "event": "opinion_parsed",
                    "task_id": task_id,
                    "opinion_index": item.opinion_index,
                    "expert_name": item.expert_name,
                    "major_category": item.major_category,
                    "content_preview": item.content[:100],
                    "progress": int(50 + (idx + 1) / total * 40),
                }

            task.expert_count = len(expert_set)
            task.opinion_count = total
            task.status = ReplyTaskStatus.GENERATING.value
            task.progress = 90
            self.db.commit()

            yield {
                "event": "done",
                "task_id": task_id,
                "progress": 100,
                "opinion_count": total,
                "expert_count": len(expert_set),
            }

        except Exception as e:
            task.status = ReplyTaskStatus.FAILED.value
            task.error_message = str(e)
            self.db.commit()
            yield {"event": "error", "task_id": task_id, "error": str(e)}

    def _pre_detect_indices(self, text: str) -> List[int]:
        """正则预检测意见编号"""
        indices = set()
        for pattern in OPINION_INDEX_PATTERNS:
            for m in pattern.finditer(text):
                try:
                    idx = int(m.group(1))
                    if 1 <= idx <= 500:
                        indices.add(idx)
                except ValueError:
                    continue
        return sorted(indices)

    async def _llm_parse_opinions(self, text: str, pre_indices: List[int]) -> List[ParsedOpinionItem]:
        """调用 LLM 解析意见文本"""
        # 如果文本过长，分段处理
        max_len = 6000
        segments = []
        if len(text) <= max_len:
            segments = [text]
        else:
            # 按编号分段
            segments = self._split_text_by_segments(text, max_len)

        all_items = []
        base_index = 0
        for seg_idx, seg in enumerate(segments):
            prompt = PARSE_USER_PROMPT.format(
                text=seg,
                pre_detected_indices=str([i - base_index for i in pre_indices if i > base_index][:30]),
            )
            messages = [
                {"role": "system", "content": PARSE_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ]
            raw = await llm_service.chat(messages, temperature=0.1, max_tokens=4096)
            items = self._parse_opinion_json(raw, base_index)
            all_items.extend(items)
            if items:
                base_index = max(it.opinion_index for it in items)

        # 去重和重新编号
        return self._deduplicate_opinions(all_items)

    def _split_text_by_segments(self, text: str, max_len: int) -> List[str]:
        """按段落切分长文本"""
        segments = []
        current = ""
        for para in re.split(r'\n\s*\n', text):
            if len(current) + len(para) > max_len and current:
                segments.append(current)
                current = para
            else:
                current = current + "\n\n" + para if current else para
        if current:
            segments.append(current)
        return segments

    def _parse_opinion_json(self, raw: str, base_index: int) -> List[ParsedOpinionItem]:
        """解析 LLM 返回的意见 JSON 数组"""
        cleaned = raw.strip()
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r'\[[\s\S]*\]', cleaned)
            if not match:
                return []
            try:
                data = json.loads(match.group(0))
            except json.JSONDecodeError:
                return []

        if not isinstance(data, list):
            return []

        items = []
        for entry in data:
            try:
                item = ParsedOpinionItem(**entry)
                items.append(item)
            except Exception:
                continue
        return items

    def _deduplicate_opinions(self, items: List[ParsedOpinionItem]) -> List[ParsedOpinionItem]:
        """去重并重新编号"""
        seen = set()
        unique = []
        for item in items:
            key = item.content[:50].strip()
            if key and key not in seen:
                seen.add(key)
                unique.append(item)
        # 重新编号
        for i, item in enumerate(unique, 1):
            item.opinion_index = i
        return unique

    # ------------------------------------------------------------------
    # 回复生成（RAG）
    # ------------------------------------------------------------------

    async def generate_reply(self, opinion_id: int, force: bool = False) -> ExpertReplyItem:
        """为单条意见生成回复"""
        opinion = self.db.query(ExpertOpinion).filter(ExpertOpinion.id == opinion_id).first()
        if not opinion:
            raise ValueError(f"意见 {opinion_id} 不存在")

        task = self.db.query(ExpertReplyTask).filter(ExpertReplyTask.id == opinion.reply_task_id).first()
        if not task:
            raise ValueError("关联任务不存在")

        # 如果已有回复且不强制重新生成，直接返回
        if opinion.reply_item and not force:
            return opinion.reply_item

        # 1. 检索报告原文上下文
        report_context = await self._retrieve_report_context(task, opinion)

        # 2. 检索规范依据
        standards_context = await self._retrieve_standards(opinion)

        # 3. LLM 生成回复
        result = await self._llm_generate_reply(opinion, report_context, standards_context)

        # 4. 保存/更新回复
        if opinion.reply_item:
            reply = opinion.reply_item
            reply.reply_content = result.reply_content
            reply.modify_status = result.modify_status
            reply.modify_location = result.modify_location
            reply.modify_page = result.modify_page
            reply.sources_json = result.sources
        else:
            reply = ExpertReplyItem(
                opinion_id=opinion.id,
                reply_content=result.reply_content,
                modify_status=result.modify_status,
                modify_location=result.modify_location,
                modify_page=result.modify_page,
                status=ReplyStatus.DRAFT.value,
                sources_json=result.sources,
            )
            self.db.add(reply)

        self.db.commit()
        self.db.refresh(reply)

        # 更新任务进度
        self._update_task_progress(task)
        return reply

    async def generate_all_replies(self, task_id: int, opinion_ids: Optional[List[int]] = None) -> AsyncGenerator[Dict[str, Any], None]:
        """批量生成回复"""
        task = self.get_task_with_details(task_id)
        if not task:
            yield {"event": "error", "error": f"任务 {task_id} 不存在"}
            return

        opinions = task.opinions
        if opinion_ids:
            opinions = [o for o in opinions if o.id in opinion_ids]

        total = len(opinions)
        for idx, opinion in enumerate(opinions):
            try:
                await self.generate_reply(opinion.id)
                yield {
                    "event": "reply_generated",
                    "task_id": task_id,
                    "opinion_id": opinion.id,
                    "opinion_index": opinion.opinion_index,
                    "progress": int((idx + 1) / total * 100),
                }
            except Exception as e:
                yield {
                    "event": "reply_error",
                    "task_id": task_id,
                    "opinion_id": opinion.id,
                    "error": str(e),
                }

        task.status = ReplyTaskStatus.EDITING.value
        self.db.commit()

        yield {
            "event": "done",
            "task_id": task_id,
            "progress": 100,
            "message": f"已生成 {total} 条回复",
        }

    def update_reply(
        self,
        opinion_id: int,
        reply_content: Optional[str] = None,
        modify_status: Optional[str] = None,
        modify_location: Optional[str] = None,
        modify_page: Optional[str] = None,
        status: Optional[str] = None,
    ) -> Optional[ExpertReplyItem]:
        """手动更新回复内容"""
        opinion = self.db.query(ExpertOpinion).filter(ExpertOpinion.id == opinion_id).first()
        if not opinion or not opinion.reply_item:
            return None
        reply = opinion.reply_item
        if reply_content is not None:
            reply.reply_content = reply_content
        if modify_status is not None:
            if modify_status not in MODIFY_STATUS_OPTIONS:
                raise ValueError(f"修改状态不合法，可选值: {MODIFY_STATUS_OPTIONS}")
            reply.modify_status = modify_status
        if modify_location is not None:
            reply.modify_location = modify_location
        if modify_page is not None:
            reply.modify_page = modify_page
        if status is not None:
            reply.status = status
        self.db.commit()
        self.db.refresh(reply)
        return reply

    # ------------------------------------------------------------------
    # RAG 检索
    # ------------------------------------------------------------------

    async def _load_document_text(self, document_id: int) -> str:
        """加载文档全文"""
        # 优先从 Chunk 表加载
        chunks = (
            self.db.query(Chunk)
            .filter(Chunk.document_id == document_id)
            .order_by(Chunk.chunk_index)
            .all()
        )
        if chunks:
            return "\n".join(
                c.chunk_text or "" for c in chunks
            )

        # 否则直接解析文件
        doc = self.db.query(Document).filter(Document.id == document_id).first()
        if not doc or not doc.file_path:
            return ""
        from app.services.document_parser import DocumentParser
        parser = DocumentParser()
        return parser.parse(doc.file_path)

    async def _retrieve_report_context(self, task: ExpertReplyTask, opinion: ExpertOpinion) -> str:
        """检索报告中与意见相关的原文片段"""
        if not task.report_document_id:
            return ""

        query_parts = [opinion.content[:200]]
        if opinion.chapter_path:
            query_parts.append(opinion.chapter_path)
        query_text = " ".join(query_parts)

        try:
            embedding = await asyncio.to_thread(embedding_service.embed_query, query_text)
            results = vector_store.search(
                embedding, top_k=5,
                filters={"document_id": task.report_document_id} if False else None,
            )
            # 过滤只取报告文档相关结果
            report_chunks = []
            chunks = (
                self.db.query(Chunk)
                .filter(Chunk.document_id == task.report_document_id)
                .all()
            )
            if chunks:
                # 简单文本匹配
                keywords = self._extract_keywords(opinion.content)
                scored = []
                for c in chunks:
                    text = c.chunk_text or ""
                    score = sum(1 for kw in keywords if kw in text)
                    if score > 0:
                        scored.append((score, text[:500]))
                scored.sort(key=lambda x: -x[0])
                report_chunks = [t for _, t in scored[:3]]

            if not report_chunks and results:
                for r in results[:3]:
                    meta = r.get("metadata", {})
                    report_chunks.append(meta.get("text", "")[:500])

            return "\n---\n".join(report_chunks) if report_chunks else ""
        except Exception as e:
            print(f"[ExpertReplyService] 报告检索失败: {e}")
            return ""

    async def _retrieve_standards(self, opinion: ExpertOpinion) -> str:
        """检索规范依据"""
        query_text = opinion.content[:300]
        if opinion.chapter_path:
            query_text = f"{opinion.chapter_path} {query_text}"

        try:
            embedding = await asyncio.to_thread(embedding_service.embed_query, query_text)
            results = vector_store.search(embedding, top_k=5)
            parts = []
            for i, r in enumerate(results, 1):
                meta = r.get("metadata", {})
                part = f"[{i}] 文件：{meta.get('file_name', '未知')}\n"
                if meta.get("page_number"):
                    part += f"页码：{meta['page_number']}\n"
                part += f"内容：{meta.get('text', '')[:400]}\n"
                parts.append(part)
            return "\n---\n".join(parts)
        except Exception as e:
            print(f"[ExpertReplyService] 规范检索失败: {e}")
            return ""

    def _extract_keywords(self, text: str, top_k: int = 8) -> List[str]:
        """提取关键词（简单的专业词典匹配）"""
        keywords = []
        for cat, kws in CATEGORY_KEYWORDS.items():
            for kw in kws:
                if kw in text:
                    keywords.append(kw)
        # 提取引号中的内容
        quoted = re.findall(r'[""「」【】《》\'\']([^""「」【】《》\'\']{2,20})', text)
        keywords.extend(quoted)
        return list(dict.fromkeys(keywords))[:top_k]

    async def _llm_generate_reply(
        self,
        opinion: ExpertOpinion,
        report_context: str,
        standards_context: str,
    ) -> ReplyGenerateResult:
        """调用 LLM 生成回复"""
        cat_name = MAJOR_CATEGORY_MAP.get(opinion.major_category or "other", "其他")

        prompt = REPLY_USER_PROMPT.format(
            opinion_index=opinion.opinion_index,
            expert_name=opinion.expert_name or "未知",
            major_category=cat_name,
            opinion_content=opinion.content,
            chapter_path=opinion.chapter_path or "未明确",
            page_number=opinion.page_number or "未明确",
            report_context=report_context[:2000] if report_context else "（未检索到报告原文）",
            standards_context=standards_context[:2000] if standards_context else "（未检索到相关规范）",
        )

        messages = [
            {"role": "system", "content": REPLY_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ]

        raw = await llm_service.chat(messages, temperature=0.3, max_tokens=2048)
        return self._parse_reply_json(raw)

    def _parse_reply_json(self, raw: str) -> ReplyGenerateResult:
        """解析回复 JSON"""
        cleaned = raw.strip()
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r'\{[\s\S]*\}', cleaned)
            if not match:
                return ReplyGenerateResult(
                    reply_content=cleaned[:500],
                    modify_status="解释说明",
                )
            try:
                data = json.loads(match.group(0))
            except json.JSONDecodeError:
                return ReplyGenerateResult(
                    reply_content=cleaned[:500],
                    modify_status="解释说明",
                )

        status = data.get("modify_status", "解释说明")
        if status not in MODIFY_STATUS_OPTIONS:
            status = "解释说明"

        return ReplyGenerateResult(
            reply_content=data.get("reply_content", ""),
            modify_status=status,
            modify_location=data.get("modify_location"),
            modify_page=data.get("modify_page"),
            sources=data.get("sources", []),
        )

    def _update_task_progress(self, task: ExpertReplyTask):
        """更新任务进度"""
        total = self.db.query(ExpertOpinion).filter(
            ExpertOpinion.reply_task_id == task.id
        ).count()
        replied = self.db.query(ExpertReplyItem).join(ExpertOpinion).filter(
            ExpertOpinion.reply_task_id == task.id
        ).count()
        task.progress = int(replied / total * 100) if total > 0 else 0
        if replied == total and total > 0:
            task.status = ReplyTaskStatus.EDITING.value
        self.db.flush()

    # ------------------------------------------------------------------
    # 导出 Word 回复表
    # ------------------------------------------------------------------

    def export_reply_table(self, task_id: int) -> Dict[str, Any]:
        """导出 Word 格式的意见回复表"""
        task = self.get_task_with_details(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")

        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx 未安装，请执行 pip install python-docx")

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("审查意见回复表")
        run.bold = True
        run.font.size = Pt(18)

        # 基本信息
        doc.add_paragraph()
        if task.meeting_name:
            doc.add_paragraph(f"审查会议：{task.meeting_name}")
        if task.meeting_date:
            doc.add_paragraph(f"审查日期：{task.meeting_date}")
        doc.add_paragraph(f"意见总数：{task.opinion_count} 条")
        doc.add_paragraph(f"专家人数：{task.expert_count} 人")
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # 回复表
        opinions_sorted = sorted(task.opinions, key=lambda x: x.order_index)
        replied_count = sum(1 for o in opinions_sorted if o.reply_item)

        table = doc.add_table(rows=1, cols=6, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "专家", "专业", "意见内容", "回复内容", "修改状态/位置"]
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True

        # 设置列宽
        widths = [Cm(1.0), Cm(1.5), Cm(1.5), Cm(5.5), Cm(6.0), Cm(2.5)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        status_color = {
            "已修改": RGBColor(0x00, 0x70, 0x00),
            "已补充": RGBColor(0x00, 0x70, 0xC0),
            "解释说明": RGBColor(0xBF, 0x8F, 0x00),
            "不修改": RGBColor(0xC0, 0x00, 0x00),
        }

        for o in opinions_sorted:
            row = table.add_row().cells
            row[0].text = str(o.opinion_index)
            row[1].text = o.expert_name or ""
            row[2].text = MAJOR_CATEGORY_MAP.get(o.major_category, o.major_category or "")
            row[3].text = o.content or ""

            if o.reply_item:
                row[4].text = o.reply_item.reply_content or ""
                status_text = o.reply_item.modify_status or ""
                loc_parts = []
                if o.reply_item.modify_location:
                    loc_parts.append(o.reply_item.modify_location)
                if o.reply_item.modify_page:
                    loc_parts.append(f"第{o.reply_item.modify_page}页")
                loc_text = "\n".join(loc_parts)
                cell_text = status_text
                if loc_text:
                    cell_text += f"\n{loc_text}"
                row[5].text = cell_text
                for p in row[5].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = status_color.get(o.reply_item.modify_status)
                        if o.reply_item.modify_status in ("已修改", "已补充"):
                            run.bold = True
            else:
                row[4].text = "（待生成回复）"
                row[5].text = "未回复"

        # 保存
        export_dir = os.path.join(settings.UPLOAD_DIR or "./uploads", "expert_replies")
        os.makedirs(export_dir, exist_ok=True)
        filename = f"审查意见回复表_任务{task_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        task.output_file_path = filepath
        task.output_filename = filename
        task.status = ReplyTaskStatus.EXPORTED.value
        self.db.commit()

        return {
            "task_id": task_id,
            "file_path": filepath,
            "file_name": filename,
            "opinion_count": len(opinions_sorted),
            "replied_count": replied_count,
            "exported_at": datetime.utcnow(),
        }


# ========================================================================
# 工厂函数
# ========================================================================

def get_expert_reply_service(db: Session) -> ExpertReplyService:
    return ExpertReplyService(db)
