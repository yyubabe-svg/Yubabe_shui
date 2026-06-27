"""
AI 初审引擎服务
- 按章节审查设计文档
- 5 个维度：param_completeness / code_compliance / chapter_completeness / value_consistency / format_standard
- 调用 vector_store 检索规范 + LLM 结构化 JSON 输出
- 问题分级：critical / major / minor / suggestion
- 导出 Word 审查意见表
"""
import os
import re
import json
import time
import asyncio
from typing import Optional, List, Dict, Any, AsyncGenerator
from datetime import datetime
from sqlalchemy.orm import Session

from app.models.ai_review import (
    AIReviewTask, AIReviewIssue,
    ReviewStatus, IssueSeverity, IssueCategory, IssueStatus,
)
from app.models.document import Document, Chunk
from app.schemas.ai_review import (
    REVIEW_DIMENSIONS, SEVERITY_ORDER, ISSUE_SEVERITY_MAP,
    ChapterIssueItem, ChapterReviewResult,
)
from app.services.llm_service import llm_service
from app.services.vector_store import vector_store
from app.services.embedding import embedding_service
from app.core.config import settings


# ========================================================================
# Prompt 模板
# ========================================================================

SYSTEM_PROMPT = """你是"蜀水智库 AI"的水利工程设计文件初审助手，服务于水利勘测设计院。
你必须严格遵守以下规则：
1. 只能基于提供的章节内容和检索到的规范依据进行审查；
2. 不得编造规范条文、数值或结论；
3. 每条问题必须引用具体规范依据，无法找到依据时标注"待人工核实"；
4. 问题严重程度严格分为四级：
   - critical：违反强制性条文，可能导致安全隐患或重大合规风险；
   - major：重要问题，影响设计质量或审批通过；
   - minor：一般问题，需修改完善；
   - suggestion：优化建议，不影响审批；
5. 输出必须是严格的 JSON，不要附加任何解释文字。
"""

CHAPTER_REVIEW_PROMPT = """请对以下设计文档章节进行审查。

【审查维度】
{dimensions_desc}

【项目信息】
项目类型提示：{project_type_hint}
设计阶段提示：{design_stage_hint}
当前章节：{chapter_path}

【章节原文】
{chapter_text}

【检索到的规范依据】
{context}

请严格按以下 JSON 结构输出审查结果（不要输出 JSON 之外的任何内容）：
{{
  "chapter_path": "{chapter_path}",
  "chapter_score": <0-100的分数>,
  "chapter_summary": "<本章节审查总结，100字以内>",
  "issues": [
    {{
      "severity": "<critical|major|minor|suggestion>",
      "category": "<param_completeness|code_compliance|chapter_completeness|value_consistency|format_standard>",
      "chapter_path": "<章节路径>",
      "page_number": <页码或null>,
      "location_desc": "<位置描述，如'第3.2节 堤顶高程计算'>",
      "description": "<问题的具体描述>",
      "basis_code": "<引用的规范名称及条文号>",
      "suggestion": "<修改建议>",
      "original_text": "<原文摘录>"
    }}
  ]
}}
"""

DIMENSION_DESC_MAP = {
    "param_completeness": "- 参数完整性：检查工程等别、建筑物级别、防洪标准、坝高/堤高、库容/流量、地震烈度、地质参数等关键设计参数是否齐全",
    "code_compliance": "- 规范符合性：核对设计内容是否符合现行规范（GB/SL/DL等），重点检查强制性条文",
    "chapter_completeness": "- 章节完整性：检查该章节应包含的内容是否齐全，是否有遗漏的必需要素",
    "value_consistency": "- 数值一致性：检查章节内及跨章节的关键数值（高程、尺寸、标准、参数等）是否前后一致",
    "format_standard": "- 格式规范性：检查编号、单位、术语、图表编号、引用格式等是否符合编制规程要求",
}


# ========================================================================
# AIReviewEngine
# ========================================================================

class AIReviewEngine:
    """AI 初审引擎"""

    def __init__(self, db: Session):
        self.db = db

    # ------------------------------------------------------------------
    # 任务管理
    # ------------------------------------------------------------------

    def create_task(
        self,
        project_id: int,
        document_id: int,
        dimensions: Optional[List[str]] = None,
        created_by: Optional[str] = None,
        project_type_hint: Optional[str] = None,
        design_stage_hint: Optional[str] = None,
    ) -> AIReviewTask:
        """创建 AI 审查任务"""
        dims = dimensions or list(REVIEW_DIMENSIONS)
        # 校验维度合法性
        invalid = [d for d in dims if d not in REVIEW_DIMENSIONS]
        if invalid:
            raise ValueError(f"不支持的审查维度: {invalid}")

        # 校验文档存在
        doc = self.db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            raise ValueError(f"文档 {document_id} 不存在")

        task = AIReviewTask(
            project_id=project_id,
            document_id=document_id,
            status=ReviewStatus.PENDING.value,
            review_dimensions=dims,
            project_type_hint=project_type_hint,
            design_stage_hint=design_stage_hint,
            progress=0,
            created_by=created_by,
        )
        self.db.add(task)
        self.db.commit()
        self.db.refresh(task)
        return task

    def get_task(self, task_id: int) -> Optional[AIReviewTask]:
        return self.db.query(AIReviewTask).filter(AIReviewTask.id == task_id).first()

    def get_task_with_issues(self, task_id: int) -> Optional[AIReviewTask]:
        from sqlalchemy.orm import joinedload
        return (
            self.db.query(AIReviewTask)
            .options(joinedload(AIReviewTask.issues))
            .filter(AIReviewTask.id == task_id)
            .first()
        )

    def list_tasks(
        self,
        project_id: Optional[int] = None,
        status: Optional[str] = None,
        page: int = 1,
        page_size: int = 20,
    ) -> List[AIReviewTask]:
        query = self.db.query(AIReviewTask)
        if project_id:
            query = query.filter(AIReviewTask.project_id == project_id)
        if status:
            query = query.filter(AIReviewTask.status == status)
        return (
            query.order_by(AIReviewTask.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
            .all()
        )

    # ------------------------------------------------------------------
    # 审查执行
    # ------------------------------------------------------------------

    async def run_review(self, task_id: int) -> AsyncGenerator[Dict[str, Any], None]:
        """
        执行审查（异步生成器，逐章节 yield 进度事件）。
        事件格式: {"event": "progress|chapter_complete|done|error", ...}
        """
        task = self.get_task(task_id)
        if not task:
            yield {"event": "error", "error": f"任务 {task_id} 不存在"}
            return

        start_time = time.time()
        task.status = ReviewStatus.REVIEWING.value
        task.progress = 0
        task.error_message = None
        self.db.commit()

        try:
            # 1. 加载文档章节
            chapters = self._load_chapters(task)
            if not chapters:
                raise ValueError("文档内容为空或无法解析章节")

            total = len(chapters)
            all_issues: List[AIReviewIssue] = []

            for idx, (chapter_path, chapter_text) in enumerate(chapters):
                task.current_chapter = chapter_path
                task.progress = int(idx / total * 100)
                self.db.commit()

                yield {
                    "event": "progress",
                    "task_id": task_id,
                    "progress": task.progress,
                    "current_chapter": chapter_path,
                }

                # 2. 检索相关规范
                retrieved = await self._retrieve_standards(
                    chapter_text=chapter_text,
                    chapter_path=chapter_path,
                    dimensions=task.review_dimensions or REVIEW_DIMENSIONS,
                )

                # 3. LLM 审查
                result = await self._review_chapter(
                    task=task,
                    chapter_path=chapter_path,
                    chapter_text=chapter_text,
                    retrieved_chunks=retrieved,
                )

                # 4. 保存问题
                chapter_issues = self._save_chapter_issues(task_id, result)
                all_issues.extend(chapter_issues)

                yield {
                    "event": "chapter_complete",
                    "task_id": task_id,
                    "current_chapter": chapter_path,
                    "chapter_result": result.model_dump(),
                    "issue_count": len(chapter_issues),
                }

            # 5. 汇总统计
            await asyncio.to_thread(self._finalize_task, task, all_issues, start_time)

            yield {
                "event": "done",
                "task_id": task_id,
                "progress": 100,
                "total_issues": len(all_issues),
                "total_score": task.total_score,
            }

        except Exception as e:
            task.status = ReviewStatus.FAILED.value
            task.error_message = str(e)
            task.completed_at = datetime.utcnow()
            task.duration_seconds = int(time.time() - start_time)
            self.db.commit()
            yield {"event": "error", "task_id": task_id, "error": str(e)}

    async def run_review_sync(self, task_id: int) -> AIReviewTask:
        """同步执行审查（收集所有事件，最终返回任务）"""
        task = self.get_task(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")

        async for _event in self.run_review(task_id):
            pass  # 消费所有事件
        self.db.refresh(task)
        return task

    # ------------------------------------------------------------------
    # 章节加载 & 切分
    # ------------------------------------------------------------------

    def _load_chapters(self, task: AIReviewTask) -> List[tuple]:
        """
        加载文档并按章节切分。
        返回 [(chapter_path, chapter_text), ...]
        优先使用 document.chapter_json，否则按标题正则切分。
        """
        doc = self.db.query(Document).filter(Document.id == task.document_id).first()
        if not doc or not doc.file_path or not os.path.exists(doc.file_path):
            return []

        # 1. 如果文档已解析有章节结构
        if doc.chapter_json:
            return self._chapters_from_json(doc.chapter_json)

        # 2. 从 Chunk 表聚合全文，然后正则切分
        chunks = (
            self.db.query(Chunk)
            .filter(Chunk.document_id == doc.id)
            .order_by(Chunk.chunk_index)
            .all()
        )
        if chunks:
            full_text = "\n".join(c.chunk_text for c in chunks if c.chunk_text)
        else:
            # 直接解析文件
            from app.services.document_parser import DocumentParser
            parser = DocumentParser()
            full_text = parser.parse(doc.file_path)

        if not full_text:
            return []

        return self._split_chapters_by_regex(full_text)

    def _chapters_from_json(self, chapter_json: Any) -> List[tuple]:
        """从 chapter_json 解析章节列表"""
        chapters = []
        if isinstance(chapter_json, list):
            for item in chapter_json:
                if isinstance(item, dict):
                    title = item.get("title", "") or item.get("chapter_path", "")
                    text = item.get("content", "") or item.get("text", "")
                    if title and text:
                        chapters.append((title, text))
                    # 递归子章节
                    for child in item.get("children", []) or []:
                        if isinstance(child, dict):
                            ct = child.get("title", "")
                            cx = child.get("content", "") or child.get("text", "")
                            if ct and cx:
                                chapters.append((ct, cx))
        return chapters

    def _split_chapters_by_regex(self, text: str) -> List[tuple]:
        """使用正则按标题层级切分章节"""
        # 匹配 "1 总则" / "1.2 设计依据" / "3.2.1 堤顶高程" 等标题
        pattern = re.compile(
            r'^(?:第[一二三四五六七八九十百]+[章节编篇]|\d+(?:\.\d+)*)[\s、．.\u3000]+([^\n]{2,60})$',
            re.MULTILINE,
        )
        matches = list(pattern.finditer(text))
        chapters = []
        if not matches:
            # 无标题结构，整体作为一章
            return [("全文", text[:8000])]

        for i, m in enumerate(matches):
            title = m.group(0).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            # 截断过长章节（避免超出 token 限制）
            if len(content) > 6000:
                content = content[:6000] + "\n...[内容过长已截断]"
            if content:
                chapters.append((title, content))

        return chapters

    # ------------------------------------------------------------------
    # 规范检索
    # ------------------------------------------------------------------

    async def _retrieve_standards(
        self,
        chapter_text: str,
        chapter_path: str,
        dimensions: List[str],
        top_k: int = 6,
    ) -> List[Dict[str, Any]]:
        """基于章节内容向量化检索相关规范条文"""
        try:
            query_text = f"{chapter_path} {chapter_text[:500]}"
            embedding = await asyncio.to_thread(
                embedding_service.embed_query, query_text
            )
            # 过滤规范类文档（file_type 包含"规范"或"标准"）
            results = vector_store.search(embedding, top_k=top_k)
            return [
                {
                    "file_name": r.get("metadata", {}).get("file_name", "未知规范"),
                    "page_number": r.get("metadata", {}).get("page_number", ""),
                    "text": r.get("metadata", {}).get("text", "")[:600],
                    "score": r.get("score", 0),
                }
                for r in results
            ]
        except Exception as e:
            print(f"[AIReviewEngine] 规范检索失败: {e}")
            return []

    # ------------------------------------------------------------------
    # LLM 章节审查
    # ------------------------------------------------------------------

    async def _review_chapter(
        self,
        task: AIReviewTask,
        chapter_path: str,
        chapter_text: str,
        retrieved_chunks: List[Dict[str, Any]],
    ) -> ChapterReviewResult:
        """调用 LLM 审查单章节，返回结构化结果"""
        dimensions_desc = "\n".join(
            DIMENSION_DESC_MAP.get(d, f"- {d}") for d in (task.review_dimensions or REVIEW_DIMENSIONS)
        )
        context = self._format_context(retrieved_chunks)

        prompt = CHAPTER_REVIEW_PROMPT.format(
            dimensions_desc=dimensions_desc,
            project_type_hint=task.project_type_hint or "未指定",
            design_stage_hint=task.design_stage_hint or "未指定",
            chapter_path=chapter_path,
            chapter_text=chapter_text,
            context=context or "（未检索到相关规范条文）",
        )

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ]

        raw = await llm_service.chat(messages, temperature=0.2, max_tokens=4096)
        return self._parse_review_json(raw, chapter_path)

    def _parse_review_json(self, raw: str, chapter_path: str) -> ChapterReviewResult:
        """解析 LLM 返回的 JSON"""
        # 清理 markdown 代码块标记
        cleaned = raw.strip()
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            # 尝试提取首个 JSON 对象
            match = re.search(r'\{[\s\S]*\}', cleaned)
            if not match:
                return ChapterReviewResult(
                    chapter_path=chapter_path,
                    issues=[],
                    chapter_summary="LLM返回解析失败",
                )
            try:
                data = json.loads(match.group(0))
            except json.JSONDecodeError:
                return ChapterReviewResult(
                    chapter_path=chapter_path,
                    issues=[],
                    chapter_summary="LLM返回解析失败",
                )

        issues = []
        for item in data.get("issues", []):
            try:
                issues.append(ChapterIssueItem(**item))
            except Exception:
                continue

        return ChapterReviewResult(
            chapter_path=data.get("chapter_path", chapter_path),
            issues=issues,
            chapter_score=data.get("chapter_score"),
            chapter_summary=data.get("chapter_summary"),
        )

    def _format_context(self, chunks: List[Dict[str, Any]]) -> str:
        parts = []
        for i, c in enumerate(chunks, 1):
            part = f"[{i}] 文件：{c.get('file_name', '未知')}\n"
            if c.get("page_number"):
                part += f"页码：{c['page_number']}\n"
            part += f"内容：{c.get('text', '')[:500]}\n"
            parts.append(part)
        return "\n---\n".join(parts)

    # ------------------------------------------------------------------
    # 问题持久化
    # ------------------------------------------------------------------

    def _save_chapter_issues(
        self, task_id: int, result: ChapterReviewResult
    ) -> List[AIReviewIssue]:
        """将单章节审查结果保存到数据库"""
        saved = []
        for item in result.issues:
            # 校验 severity
            severity = item.severity if item.severity in SEVERITY_ORDER else "minor"
            category = item.category if item.category in REVIEW_DIMENSIONS else "format_standard"

            issue = AIReviewIssue(
                review_task_id=task_id,
                severity=severity,
                category=category,
                chapter_path=item.chapter_path or result.chapter_path,
                page_number=item.page_number,
                location_desc=item.location_desc,
                description=item.description,
                basis_code=item.basis_code,
                suggestion=item.suggestion,
                original_text=item.original_text,
                status=IssueStatus.OPEN.value,
                severity_order=SEVERITY_ORDER.get(severity, 2),
            )
            self.db.add(issue)
            saved.append(issue)
        self.db.flush()
        return saved

    def _finalize_task(
        self, task: AIReviewTask, all_issues: List[AIReviewIssue], start_time: float
    ):
        """汇总问题统计，计算总分，完成任务"""
        critical = sum(1 for i in all_issues if i.severity == "critical")
        major = sum(1 for i in all_issues if i.severity == "major")
        minor = sum(1 for i in all_issues if i.severity == "minor")
        suggestion = sum(1 for i in all_issues if i.severity == "suggestion")

        # 简单评分：满分100，critical 扣20，major 扣10，minor 扣3，suggestion 扣1
        score = 100 - critical * 20 - major * 10 - minor * 3 - suggestion * 1
        score = max(0, min(100, score))

        task.issue_count_critical = critical
        task.issue_count_major = major
        task.issue_count_minor = minor
        task.issue_count_suggestion = suggestion
        task.total_score = float(score)
        task.summary = (
            f"审查完成，共发现 {len(all_issues)} 个问题："
            f"严重 {critical} 条，重要 {major} 条，一般 {minor} 条，建议 {suggestion} 条。"
        )
        task.progress = 100
        task.current_chapter = None
        task.status = ReviewStatus.COMPLETED.value
        task.completed_at = datetime.utcnow()
        task.duration_seconds = int(time.time() - start_time)
        self.db.commit()

    # ------------------------------------------------------------------
    # 问题状态更新
    # ------------------------------------------------------------------

    def update_issue(self, issue_id: int, status: Optional[str] = None, note: Optional[str] = None) -> Optional[AIReviewIssue]:
        issue = self.db.query(AIReviewIssue).filter(AIReviewIssue.id == issue_id).first()
        if not issue:
            return None
        if status:
            issue.status = status
        if note is not None:
            issue.note = note
        self.db.commit()
        self.db.refresh(issue)
        return issue

    # ------------------------------------------------------------------
    # 导出 Word 审查意见表
    # ------------------------------------------------------------------

    def export_report(self, task_id: int) -> Dict[str, Any]:
        """导出 Word 审查意见表，返回文件信息"""
        task = self.get_task_with_issues(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")
        if task.status != ReviewStatus.COMPLETED.value:
            raise ValueError("任务尚未完成，无法导出")

        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx 未安装，请执行 pip install python-docx")

        doc = DocxDocument()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)

        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("AI 初审意见表")
        run.bold = True
        run.font.size = Pt(18)

        # 基本信息表
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=4, style="Table Grid")
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("项目ID", str(task.project_id), "文档ID", str(task.document_id)),
            ("审查时间", task.completed_at.strftime("%Y-%m-%d %H:%M") if task.completed_at else "", "耗时", f"{task.duration_seconds or 0}秒"),
            ("严重问题", str(task.issue_count_critical), "重要问题", str(task.issue_count_major)),
            ("一般问题", str(task.issue_count_minor), "建议优化", str(task.issue_count_suggestion)),
            ("综合评分", f"{task.total_score:.1f}/100", "审查人", task.created_by or "AI"),
        ]
        for r, row_data in enumerate(info_data):
            for c, val in enumerate(row_data):
                cell = info_table.cell(r, c)
                cell.text = val
                if c % 2 == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

        # 审查总结
        doc.add_paragraph()
        summary_heading = doc.add_paragraph()
        run = summary_heading.add_run("一、审查总结")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(task.summary or "")

        # 问题明细表
        doc.add_paragraph()
        issue_heading = doc.add_paragraph()
        run = issue_heading.add_run("二、问题明细")
        run.bold = True
        run.font.size = Pt(14)

        # 按严重程度排序
        issues_sorted = sorted(task.issues, key=lambda x: (x.severity_order or 2, x.id))

        issue_table = doc.add_table(rows=1, cols=8, style="Table Grid")
        issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "严重程度", "维度", "章节", "页码", "问题描述", "规范依据", "修改建议"]
        hdr_cells = issue_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True

        severity_color = {
            "critical": RGBColor(0xC0, 0x00, 0x00),
            "major": RGBColor(0xED, 0x7D, 0x31),
            "minor": RGBColor(0xBF, 0xBF, 0x00),
            "suggestion": RGBColor(0x00, 0x70, 0xC0),
        }

        for idx, issue in enumerate(issues_sorted, 1):
            row_cells = issue_table.add_row().cells
            row_cells[0].text = str(idx)
            sev_label = ISSUE_SEVERITY_MAP.get(issue.severity, issue.severity)
            row_cells[1].text = sev_label
            # 设置严重程度颜色
            for p in row_cells[1].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = severity_color.get(issue.severity)
                    run.bold = True
            row_cells[2].text = issue.category or ""
            row_cells[3].text = issue.chapter_path or ""
            row_cells[4].text = str(issue.page_number) if issue.page_number else ""
            row_cells[5].text = issue.description or ""
            row_cells[6].text = issue.basis_code or ""
            row_cells[7].text = issue.suggestion or ""

        # 保存文件
        export_dir = os.path.join(settings.UPLOAD_DIR or "./uploads", "ai_review_reports")
        os.makedirs(export_dir, exist_ok=True)
        filename = f"AI初审意见表_任务{task_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        task.output_file_path = filepath  # 复用字段
        self.db.commit()

        return {
            "task_id": task_id,
            "file_path": filepath,
            "file_name": filename,
            "issue_count": len(issues_sorted),
            "exported_at": datetime.utcnow(),
        }


# ========================================================================
# 工厂函数
# ========================================================================

def get_ai_review_engine(db: Session) -> AIReviewEngine:
    return AIReviewEngine(db)
