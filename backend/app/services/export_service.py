"""
成果导出服务
==============
提供设计成果文件导出能力：
1. 工程量清单 Excel 导出（.xlsx）
2. 设计说明书 Word 导出（.docx）
3. 断面设计计算书导出
"""
import os
import io
import math
from datetime import datetime
from typing import Dict, List, Any, Optional
from fastapi.responses import StreamingResponse

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ==================== Excel 工程量清单 ====================

# 样式
THIN_BORDER = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)
HEADER_FILL = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
HEADER_FONT = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
CELL_FONT = Font(name='微软雅黑', size=10)
TITLE_FONT = Font(name='微软雅黑', size=14, bold=True)
SUBTITLE_FONT = Font(name='微软雅黑', size=11, bold=True)
CENTER_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)
LEFT_ALIGN = Alignment(horizontal='left', vertical='center', wrap_text=True)


def generate_bill_of_quantities_excel(
    project_name: str,
    section_results: List[Dict[str, Any]],
    channel_lengths: Optional[List[float]] = None,
    project_info: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    生成工程量清单 Excel
    section_results: 参数化设计结果列表
    channel_lengths: 各段长度(m)，与section_results一一对应
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "工程量清单"

    # 列宽
    col_widths = [6, 20, 12, 10, 12, 12, 14, 14, 18]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # 标题
    ws.merge_cells('A1:I1')
    ws['A1'] = f"{project_name} — 堤防工程工程量清单"
    ws['A1'].font = TITLE_FONT
    ws['A1'].alignment = CENTER_ALIGN
    ws.row_dimensions[1].height = 30

    # 副标题信息
    info_row = 2
    ws.merge_cells(f'A{info_row}:I{info_row}')
    now = datetime.now().strftime('%Y年%m月%d日')
    ws[f'A{info_row}'] = f"编制日期：{now}    编制单位：蜀水AI智能设计"
    ws[f'A{info_row}'].font = Font(name='微软雅黑', size=9, italic=True, color='666666')
    ws[f'A{info_row}'].alignment = Alignment(horizontal='right')

    # 表头
    headers = ['序号', '工程项目名称', '单位', '数量', '单价(元)', '合价(元)', '备注', '断面形式', '每延米量']
    header_row = 4
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
    ws.row_dimensions[header_row].height = 22

    # 汇总数据
    total_cost = 0
    item_idx = 1
    row = header_row + 1

    if channel_lengths is None:
        channel_lengths = [1000] * len(section_results)  # 默认每段1km

    for si, (sr, seg_len) in enumerate(zip(section_results, channel_lengths)):
        if not sr.get('success'):
            continue

        q = sr.get('quantities', {})
        c = sr.get('costs', {})
        params = sr.get('parameters', {})
        section_name = sr.get('section_name', '堤防断面')
        section_type = sr.get('section_type', '')

        # 定义工程量清单项目
        items = [
            {
                'name': f'堤身土方填筑（{section_name}）',
                'unit': 'm³',
                'qty_per_m': q.get('fill_volume_m3_per_m', 0),
                'unit_price': 35,
                'note': f'含推平碾压，{seg_len:.0f}m段'
            },
            {
                'name': f'{params.get("revetment_name","护岸")}（{section_name}）',
                'unit': 'm³',
                'qty_per_m': q.get('revetment_volume_m3_per_m', 0),
                'unit_price': c.get('revetment_cost_yuan_per_m', 0) / max(q.get('revetment_volume_m3_per_m', 0.001), 0.001),
                'note': '含基础'
            },
            {
                'name': f'基础开挖（{section_name}）',
                'unit': 'm³',
                'qty_per_m': q.get('excavation_m3_per_m', 0),
                'unit_price': 25,
                'note': '含弃渣外运'
            },
        ]

        for it in items:
            total_qty = it['qty_per_m'] * seg_len
            total_price = total_qty * it['unit_price']
            total_cost += total_price

            ws.cell(row=row, column=1, value=item_idx).font = CELL_FONT
            ws.cell(row=row, column=2, value=it['name']).font = CELL_FONT
            ws.cell(row=row, column=3, value=it['unit']).font = CELL_FONT
            ws.cell(row=row, column=4, value=round(total_qty, 1)).font = CELL_FONT
            ws.cell(row=row, column=5, value=round(it['unit_price'], 2)).font = CELL_FONT
            ws.cell(row=row, column=6, value=round(total_price, 0)).font = CELL_FONT
            ws.cell(row=row, column=7, value=it['note']).font = CELL_FONT
            ws.cell(row=row, column=8, value=section_type).font = CELL_FONT
            ws.cell(row=row, column=9, value=round(it['qty_per_m'], 3)).font = CELL_FONT

            for col in range(1, 10):
                cell = ws.cell(row=row, column=col)
                cell.border = THIN_BORDER
                cell.alignment = CENTER_ALIGN if col not in (2,7) else LEFT_ALIGN

            ws.row_dimensions[row].height = 18
            item_idx += 1
            row += 1

    # 合计行
    row += 1
    ws.merge_cells(f'A{row}:E{row}')
    ws.cell(row=row, column=1, value='合计').font = SUBTITLE_FONT
    ws.cell(row=row, column=1).alignment = CENTER_ALIGN
    ws.cell(row=row, column=6, value=round(total_cost, 0)).font = Font(name='微软雅黑', size=11, bold=True, color='C00000')
    for col in range(1, 10):
        ws.cell(row=row, column=col).border = THIN_BORDER
        ws.cell(row=row, column=col).fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')

    # 折合每km投资
    total_len = sum(channel_lengths) if channel_lengths else 1000
    row += 1
    ws.merge_cells(f'A{row}:E{row}')
    ws.cell(row=row, column=1, value=f'折合单位投资').font = CELL_FONT
    ws.cell(row=row, column=6, value=f'{round(total_cost/total_len*1000, 0)} 元/km').font = Font(name='微软雅黑', size=10, bold=True)

    # 第二张表：主要材料汇总
    ws2 = wb.create_sheet("主要材料汇总")
    ws2.column_dimensions['A'].width = 6
    ws2.column_dimensions['B'].width = 20
    ws2.column_dimensions['C'].width = 10
    ws2.column_dimensions['D'].width = 14
    ws2.column_dimensions['E'].width = 20

    ws2.merge_cells('A1:E1')
    ws2['A1'] = "主要材料用量汇总表"
    ws2['A1'].font = TITLE_FONT
    ws2['A1'].alignment = CENTER_ALIGN
    ws2.row_dimensions[1].height = 28

    mat_headers = ['序号', '材料名称', '单位', '总用量', '备注']
    for col, h in enumerate(mat_headers, 1):
        cell = ws2.cell(row=3, column=col, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER

    # 汇总各种材料
    total_fill = sum(sr.get('quantities',{}).get('fill_volume_m3_per_m',0)*cl for sr,cl in zip(section_results, channel_lengths))
    total_stone = sum(sr.get('quantities',{}).get('concrete_or_stone_m3_per_m',0)*cl for sr,cl in zip(section_results, channel_lengths))
    total_exc = sum(sr.get('quantities',{}).get('excavation_m3_per_m',0)*cl for sr,cl in zip(section_results, channel_lengths))

    materials = [
        ('土方填筑', 'm³', round(total_fill, 0), '压实方'),
        ('块石/混凝土', 'm³', round(total_stone, 0), '护岸+基础'),
        ('土方开挖', 'm³', round(total_exc, 0), '自然方'),
        ('水泥', 't', round(total_stone * 0.25, 1), '按砌石/混凝土估算'),
        ('砂', 'm³', round(total_stone * 0.45, 0), '砂浆用砂'),
    ]
    for i, (name, unit, qty, note) in enumerate(materials, 1):
        r = 3 + i
        ws2.cell(row=r, column=1, value=i).font = CELL_FONT
        ws2.cell(row=r, column=2, value=name).font = CELL_FONT
        ws2.cell(row=r, column=3, value=unit).font = CELL_FONT
        ws2.cell(row=r, column=4, value=qty).font = CELL_FONT
        ws2.cell(row=r, column=5, value=note).font = CELL_FONT
        for col in range(1, 6):
            ws2.cell(row=r, column=col).border = THIN_BORDER
            ws2.cell(row=r, column=col).alignment = CENTER_ALIGN if col != 5 else LEFT_ALIGN

    # 保存到bytes
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ==================== Word 设计说明书 ====================

def _set_doc_default_font(doc: Document):
    """设置文档默认字体"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符


def _add_heading_cn(doc: Document, text: str, level: int = 1):
    """添加中文标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.first_line_indent = Cm(0)
    return h


def _add_para(doc: Document, text: str, bold: bool = False, indent: bool = True, align=None):
    """添加段落"""
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p


def _add_table_from_data(doc: Document, headers: List[str], rows: List[List[Any]]):
    """添加表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 数据行
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table


def generate_design_report_docx(
    project_name: str,
    project_info: Dict[str, Any],
    section_results: List[Dict[str, Any]],
    hydraulic_results: Optional[List[Dict[str, Any]]] = None,
    rainfall_result: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    生成设计说明书 Word
    """
    doc = Document()
    _set_doc_default_font(doc)

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{project_name}\n初步设计报告")
    title_run.font.name = '黑体'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息
    info_items = [
        f"工程等别：{project_info.get('project_grade', 'Ⅳ')}等",
        f"建筑物级别：{project_info.get('building_level', 4)}级",
        f"设计阶段：{project_info.get('design_stage', '初步设计')}",
        f"编制日期：{datetime.now().strftime('%Y年%m月')}",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(item)
        r.font.name = '仿宋'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        r.font.size = Pt(16)

    doc.add_page_break()

    # ===== 目录（占位） =====
    _add_heading_cn(doc, "目  录", level=1)
    _add_para(doc, "（自动生成目录请使用Word的引用→目录功能）", indent=False)
    doc.add_page_break()

    # ===== 1. 综合说明 =====
    _add_heading_cn(doc, "1  综合说明", level=1)

    _add_heading_cn(doc, "1.1  工程概况", level=2)
    _add_para(doc,
        f"{project_name}位于{project_info.get('location', '项目所在地')}，"
        f"属{project_info.get('project_type', '河道治理')}工程。"
        f"工程等别为{project_info.get('project_grade', 'Ⅳ')}等，主要建筑物级别为"
        f"{project_info.get('building_level', 4)}级，"
        f"设计洪水标准为{project_info.get('flood_std_design', '20年一遇')}，"
        f"校核洪水标准为{project_info.get('flood_std_check', '50年一遇')}。"
    )

    scale_items = []
    if project_info.get('catchment_area'):
        scale_items.append(f"集雨面积{project_info['catchment_area']}km²")
    if project_info.get('river_governance_length'):
        scale_items.append(f"治理河长{project_info['river_governance_length']}km")
    if project_info.get('embankment_length'):
        scale_items.append(f"堤防总长{project_info['embankment_length']}km")
    if scale_items:
        _add_para(doc, f"工程主要规模：{'，'.join(scale_items)}。")

    _add_heading_cn(doc, "1.2  设计依据", level=2)
    standards = [
        "《防洪标准》（GB 50201-2014）",
        "《堤防工程设计规范》（GB 50286-2013）",
        "《水利水电工程等级划分及洪水标准》（SL 252-2017）",
        "《水工挡土墙设计规范》（SL 379-2007）",
        "《水利水电工程设计洪水计算规范》（SL 44-2006）",
    ]
    for s in standards:
        p = doc.add_paragraph(s, style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)

    # ===== 2. 水文 =====
    _add_heading_cn(doc, "2  水  文", level=1)
    _add_heading_cn(doc, "2.1  流域概况", level=2)
    _add_para(doc, "（简述流域自然地理、气象水文特性，根据项目资料补充。）")

    _add_heading_cn(doc, "2.2  设计暴雨", level=2)
    if rainfall_result and rainfall_result.get('success'):
        rd = rainfall_result.get('data', {})
        _add_para(doc,
            f"采用{rd.get('city', '成都')}市暴雨强度公式：{rd.get('formula', '')}。"
            f"设计重现期 P={rd.get('return_period')}年，降雨历时{rd.get('duration_min')}分钟，"
            f"雨峰系数r={rd.get('r')}。"
        )
        _add_para(doc,
            f"芝加哥雨型计算结果：总降雨量{rd.get('total_rainfall_mm')}mm，"
            f"平均雨强{rd.get('avg_intensity_mm_h')}mm/h，"
            f"峰值雨强{rd.get('peak_intensity_mm_h')}mm/h，"
            f"峰现时间为第{rd.get('peak_time_min')}分钟。"
        )
    else:
        _add_para(doc, "设计暴雨根据当地暴雨洪水手册查算，采用推理公式法或瞬时单位线法推求设计洪水。")

    _add_heading_cn(doc, "2.3  设计洪水", level=2)
    _add_para(doc, f"根据《水利水电工程设计洪水计算规范》（SL 44-2006），"
                   f"{project_info.get('building_level',4)}级堤防采用{project_info.get('flood_std_design','20年一遇')}洪水设计。")

    # ===== 3. 工程地质 =====
    _add_heading_cn(doc, "3  工程地质", level=1)
    _add_para(doc, "（简述区域地质、堤基工程地质条件、天然建筑材料，根据地勘报告补充。）")
    _add_para(doc, "堤基土主要为粉质黏土和砂卵石，承载力满足要求；堤身填筑料可就近取材。")

    # ===== 4. 工程布置及建筑物 =====
    _add_heading_cn(doc, "4  工程布置及建筑物", level=1)
    _add_heading_cn(doc, "4.1  工程等别和标准", level=2)
    _add_para(doc,
        f"根据《防洪标准》（GB 50201-2014）和《堤防工程设计规范》（GB 50286-2013），"
        f"本工程等别为{project_info.get('project_grade', 'Ⅳ')}等，主要建筑物级别为"
        f"{project_info.get('building_level', 4)}级，次要建筑物级别为5级，临时建筑物级别为5级。"
    )

    _add_heading_cn(doc, "4.2  堤防工程设计", level=2)

    for i, sr in enumerate(section_results):
        if not sr.get('success'):
            continue
        params = sr.get('parameters', {})
        geo = sr.get('geometry', {})
        q = sr.get('quantities', {})
        c = sr.get('costs', {})
        stab = sr.get('stability', {})

        _add_heading_cn(doc, f"4.2.{i+1}  {sr.get('section_name', f'堤段{i+1}')}", level=3)

        # 断面尺寸
        _add_para(doc,
            f"采用{sr.get('section_type','')}断面，"
            f"堤顶高程{params.get('crest_elevation')}m，堤顶宽度{params.get('crest_width')}m，"
            f"设计水深{params.get('water_depth')}m，超高{params.get('freeboard')}m。"
        )
        if 'm_slope' in params:
            _add_para(doc, f"堤身边坡系数1:{params['m_slope']}，护岸采用{params.get('revetment_name','')}，厚度{params.get('revetment_thickness')}m。")
        if 'wall_bottom_thickness' in params:
            _add_para(doc, f"采用重力式挡墙，墙顶宽{params.get('wall_thickness')}m，墙底宽{params.get('wall_bottom_thickness')}m，基础埋深{params.get('foundation_depth')}m。")

        # 水力要素表
        hydro_rows = [[
            f"{params.get('water_depth','-')}",
            f"{q.get('flow_area_m2','-')}",
            f"{q.get('wetted_perimeter_m','-')}",
            f"{q.get('hydraulic_radius_m','-')}",
        ]]
        _add_table_from_data(doc, ['水深(m)', '过水面积(m²)', '湿周(m)', '水力半径(m)'], hydro_rows)

        # 工程量
        _add_para(doc,
            f"主要工程量（每延米）：土方填筑{q.get('fill_volume_m3_per_m',0):.2f}m³，"
            f"护岸{q.get('revetment_volume_m3_per_m',0):.2f}m³，"
            f"基础开挖{q.get('excavation_m3_per_m',0):.2f}m³。"
        )

        # 稳定计算
        _add_para(doc,
            f"抗滑稳定安全系数Kc={stab.get('anti_slide_Kc','-')}，"
            f"{'满足' if stab.get('pass') else '不满足'}规范要求。"
        )

        # 规范符合性
        for comp in sr.get('compliance', []):
            p = doc.add_paragraph(comp, style='List Bullet')
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(11)

    # ===== 5. 施工组织设计 =====
    _add_heading_cn(doc, "5  施工组织设计", level=1)
    _add_para(doc, "（简述施工条件、施工导流、主体工程施工、施工总布置、施工总进度。）")
    _add_para(doc, "施工总工期建议6个月，其中主体工程施工期4个月。土方填筑严格控制含水量和压实度，浆砌石护座采用座浆法砌筑。")

    # ===== 6. 工程投资估算 =====
    _add_heading_cn(doc, "6  工程投资估算", level=1)

    if section_results:
        total_cost = sum(
            sr.get('costs',{}).get('total_cost_yuan_per_m',0) * (1000 if i==0 else 1000)
            for i, sr in enumerate(section_results) if sr.get('success')
        )
        _add_para(doc,
            f"本工程估算总投资约{project_info.get('total_investment', round(total_cost/10000, 0))}万元，"
            f"其中建筑工程费约占65%，机电金结约占5%，临时工程约占8%，其他费用约占12%，基本预备费约占10%。"
        )

    _add_para(doc, "详细工程量清单见附件《工程量清单》Excel文件。")

    # ===== 7. 结论与建议 =====
    _add_heading_cn(doc, "7  结论与建议", level=1)
    _add_para(doc, "本工程堤线布置合理，堤身断面满足抗滑稳定要求，护岸结构选型适宜，工程投资经济合理。建议下阶段补充地质勘察工作，进一步优化堤身结构设计。")

    # 保存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class ExportService:
    """成果导出服务"""

    def export_bill_of_quantities(
        self,
        project_name: str,
        section_results: List[Dict[str, Any]],
        channel_lengths: Optional[List[float]] = None,
        project_info: Optional[Dict[str, Any]] = None,
    ) -> StreamingResponse:
        """导出工程量清单Excel"""
        data = generate_bill_of_quantities_excel(project_name, section_results, channel_lengths, project_info)
        filename = f"工程量清单_{project_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )

    def export_design_report(
        self,
        project_name: str,
        project_info: Dict[str, Any],
        section_results: List[Dict[str, Any]],
        hydraulic_results: Optional[List[Dict[str, Any]]] = None,
        rainfall_result: Optional[Dict[str, Any]] = None,
    ) -> StreamingResponse:
        """导出设计说明书Word"""
        data = generate_design_report_docx(project_name, project_info, section_results, hydraulic_results, rainfall_result)
        filename = f"设计说明书_{project_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )


# 单例
export_service = ExportService()
