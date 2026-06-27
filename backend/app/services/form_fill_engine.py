"""通用表格填报引擎

核心能力：
1. create_task  - 创建填报任务（关联项目、模板、源文档）
2. extract_fields - 调用 LLM 从源文档提取字段值
3. fill_template  - 将确认后的数据填充到 Word 模板，支持三种定位方式：
   - table_cell(table_idx, row, col)
   - bookmark(bookmark_name)
   - label_adjacent(label_text, direction)
4. 复选框处理（□/■/☑ 替换）、保留单元格原有字体格式
"""
import os
import re
import json
import shutil
import copy
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from app.core.config import settings
from app.core.database import SessionLocal
from app.models.form_template import (
    FormFillTask, FormField, FormFillFieldValue,
    FormFillStatus, FormFieldType, LocatorType,
)
from app.models.document import Document as DocumentModel
from app.models.form_template import FormTemplate
from app.services.llm_service import llm_service
from app.services.document_parser import document_parser


# ------------------------------------------------------------------
# LLM 提取 Prompt
# ------------------------------------------------------------------

FIELD_EXTRACT_PROMPT = """你是水利工程设计文件信息提取专家。请根据下面的表单字段说明，从项目设计报告中提取对应的值。

表单字段列表（JSON数组，每项包含 field_key、field_label、field_type、llm_extract_hint）：
{fields_json}

项目报告内容：
{report_content}

请严格按照 JSON 格式输出，结构如下：
{{
  "fields": [
    {{
      "field_key": "字段标识",
      "value": "提取的值（字符串），无法确定则填 null",
      "confidence": 0.0~1.0 的置信度,
      "source_text": "原文摘录片段（便于溯源）",
      "source_section": "来源章节（如能判断）"
    }}
  ]
}}

注意：
- 只输出 JSON，不要添加任何说明文字。
- checkbox 类型字段，若判断为"是/有/勾选"则 value 填 "true"，否则填 "false"。
- date 类型字段尽量输出 YYYY年M月 或 YYYY-MM-DD 格式。
- number 类型字段只输出数字（可带单位）。
"""


# ------------------------------------------------------------------
# FormFillEngine
# ------------------------------------------------------------------

class FormFillEngine:
    """通用表格填报引擎"""

    # 输出目录
    OUTPUT_SUBDIR = "form_fill_output"

    # ------------------------------------------------------------------
    # 公开 API
    # ------------------------------------------------------------------

    def create_task(
        self,
        project_id: int,
        template_id: int,
        document_id: int,
        created_by: Optional[str] = None,
    ) -> FormFillTask:
        """创建填报任务。

        :returns: 持久化后的 FormFillTask 实例
        """
        db = SessionLocal()
        try:
            # 校验引用存在
            template = db.query(FormTemplate).filter(FormTemplate.id == template_id).first()
            if not template:
                raise ValueError(f"模板不存在: template_id={template_id}")
            doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
            if not doc:
                raise ValueError(f"文档不存在: document_id={document_id}")

            task = FormFillTask(
                project_id=project_id,
                template_id=template_id,
                document_id=document_id,
                status=FormFillStatus.PENDING.value,
                progress=0,
                created_by=created_by,
                created_at=datetime.utcnow(),
            )
            db.add(task)
            db.commit()
            db.refresh(task)
            return task
        finally:
            db.close()

    async def extract_fields(self, task_id: int) -> Dict[str, Any]:
        """读取源文档 -> LLM 提取字段 -> 写入 FormFillFieldValue。

        :returns: {"task_id", "status", "fields": [...], "warnings": [...]}
        """
        db = SessionLocal()
        try:
            task = db.query(FormFillTask).filter(FormFillTask.id == task_id).first()
            if not task:
                raise ValueError(f"任务不存在: task_id={task_id}")

            # 更新状态
            task.status = FormFillStatus.EXTRACTING.value
            task.progress = 10
            db.commit()

            # 1. 获取模板字段定义
            fields = db.query(FormField).filter(
                FormField.template_id == task.template_id
            ).order_by(FormField.sort_order, FormField.id).all()
            if not fields:
                task.status = FormFillStatus.FAILED.value
                task.error_message = "模板未配置任何字段"
                db.commit()
                return {
                    "task_id": task_id,
                    "status": FormFillStatus.FAILED.value,
                    "fields": [],
                    "warnings": ["模板未配置字段"],
                }

            # 2. 读取源文档文本
            doc_model = db.query(DocumentModel).filter(
                DocumentModel.id == task.document_id
            ).first()
            text_content = ""
            if doc_model and doc_model.file_path and os.path.exists(doc_model.file_path):
                text_content = document_parser.parse(doc_model.file_path)
            task.progress = 30
            db.commit()

            if not text_content:
                task.status = FormFillStatus.FAILED.value
                task.error_message = "源文档解析为空或文件不存在"
                db.commit()
                return {
                    "task_id": task_id,
                    "status": FormFillStatus.FAILED.value,
                    "fields": [],
                    "warnings": ["源文档解析失败或内容为空"],
                }

            # 3. 组装字段描述传给 LLM
            fields_desc = []
            for f in fields:
                fields_desc.append({
                    "field_key": f.field_key,
                    "field_label": f.field_label,
                    "field_type": f.field_type,
                    "llm_extract_hint": f.llm_extract_hint or "",
                })

            # 截断过长文本（保护上下文窗口）
            truncated = text_content[:12000] if len(text_content) > 12000 else text_content
            prompt = FIELD_EXTRACT_PROMPT.format(
                fields_json=json.dumps(fields_desc, ensure_ascii=False, indent=2),
                report_content=truncated,
            )

            task.progress = 50
            db.commit()

            # 4. 调用 LLM
            extracted_map: Dict[str, Dict[str, Any]] = {}
            warnings: List[str] = []
            try:
                llm_response = await llm_service.generate(prompt)
                json_match = re.search(r'\{[\s\S]*\}', llm_response)
                if json_match:
                    data = json.loads(json_match.group())
                    for item in data.get("fields", []):
                        fk = item.get("field_key")
                        if fk:
                            extracted_map[fk] = item
                else:
                    warnings.append("LLM 返回内容非 JSON 格式，提取失败")
            except Exception as e:
                warnings.append(f"LLM 提取异常: {e}")

            # 5. 写入 FormFillFieldValue
            field_values_result = []
            for f in fields:
                item = extracted_map.get(f.field_key, {})
                raw_value = item.get("value")
                confidence = float(item.get("confidence", 0.0) or 0.0)
                source_text = item.get("source_text")
                source_section = item.get("source_section")

                # 处理默认值
                if (raw_value is None or raw_value == "") and f.default_value:
                    raw_value = f.default_value
                    if confidence < 0.3:
                        confidence = 0.5  # 默认值给一个中等置信度

                value_str = str(raw_value) if raw_value is not None else None

                fv = db.query(FormFillFieldValue).filter(
                    FormFillFieldValue.task_id == task_id,
                    FormFillFieldValue.field_id == f.id,
                ).first()
                if not fv:
                    fv = FormFillFieldValue(
                        task_id=task_id,
                        field_id=f.id,
                    )
                    db.add(fv)
                fv.extracted_value = value_str
                fv.confidence = confidence
                fv.source_text = source_text
                fv.source_section = source_section
                # 初次提取时 confirmed_value 暂存与 extracted_value 相同（用户可修改）
                fv.confirmed_value = value_str
                fv.updated_at = datetime.utcnow()

                field_values_result.append({
                    "field_key": f.field_key,
                    "value": value_str,
                    "confidence": confidence,
                    "source_text": source_text,
                    "source_section": source_section,
                })

                if f.required and (value_str is None or value_str == ""):
                    warnings.append(f"必填字段「{f.field_label}」未能自动提取，请人工补充")

            # 6. 更新任务状态
            task.extracted_data_json = field_values_result
            task.status = FormFillStatus.AWAITING_CONFIRMATION.value
            task.progress = 70
            task.updated_at = datetime.utcnow()
            db.commit()

            return {
                "task_id": task_id,
                "status": task.status,
                "fields": field_values_result,
                "warnings": warnings,
                "message": "字段提取完成，请确认后进行模板填充",
            }
        finally:
            db.close()

    def fill_template(
        self,
        task_id: int,
        confirmed_data: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """根据确认后的数据填充 Word 模板并生成输出文件。

        :param task_id: 任务 ID
        :param confirmed_data: {field_key: value, ...}，
               若为 None 则自动从 DB 中读取每个字段的 confirmed_value。
               checkbox 字段 value 可为 bool 或 "true"/"false" 字符串。
        :returns: {"task_id", "status", "output_file", "filled_count", "warnings": [...]}
        """
        db = SessionLocal()
        try:
            task = db.query(FormFillTask).filter(FormFillTask.id == task_id).first()
            if not task:
                raise ValueError(f"任务不存在: task_id={task_id}")

            template = db.query(FormTemplate).filter(FormTemplate.id == task.template_id).first()
            if not template or not template.template_file_path or not os.path.exists(template.template_file_path):
                raise FileNotFoundError(f"模板文件不存在: {getattr(template, 'template_file_path', None)}")

            task.status = FormFillStatus.FILLING.value
            task.progress = 75
            db.commit()

            # 1. 准备字段映射: field_key -> (FormField, value)
            fields = db.query(FormField).filter(
                FormField.template_id == task.template_id
            ).all()
            field_map: Dict[str, Tuple[FormField, Any]] = {}
            for f in fields:
                if confirmed_data and f.field_key in confirmed_data:
                    val = confirmed_data[f.field_key]
                else:
                    # 回退到 DB 中的 confirmed_value
                    fv = db.query(FormFillFieldValue).filter(
                        FormFillFieldValue.task_id == task_id,
                        FormFillFieldValue.field_id == f.id,
                    ).first()
                    val = fv.confirmed_value if fv else None
                field_map[f.field_key] = (f, val)

                # 更新 confirmed_value
                if confirmed_data and f.field_key in confirmed_data:
                    fv = db.query(FormFillFieldValue).filter(
                        FormFillFieldValue.task_id == task_id,
                        FormFillFieldValue.field_id == f.id,
                    ).first()
                    if fv:
                        fv.confirmed_value = str(val) if val is not None else None
                        fv.updated_at = datetime.utcnow()

            # 保存 confirmed_data 到任务
            if confirmed_data:
                task.confirmed_data_json = confirmed_data

            # 2. 复制模板到输出路径
            output_dir = os.path.join(settings.UPLOAD_DIR, self.OUTPUT_SUBDIR)
            os.makedirs(output_dir, exist_ok=True)

            base_name = os.path.splitext(os.path.basename(template.template_file_path))[0]
            safe_name = re.sub(r'[<>:"/\\|?*]', '_', base_name)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{safe_name}_filled_{ts}.docx"
            output_path = os.path.join(output_dir, output_filename)
            shutil.copy2(template.template_file_path, output_path)

            # 3. 打开 docx 并填充
            doc = Document(output_path)
            filled_count = 0
            warnings: List[str] = []

            for field_key, (field_def, value) in field_map.items():
                if value is None or (isinstance(value, str) and value.strip() == ""):
                    if field_def.required:
                        warnings.append(f"必填字段「{field_def.field_label}」({field_key}) 值为空，跳过填充")
                    continue

                locator_type = field_def.locator_type or LocatorType.TABLE_CELL.value
                locator_data = field_def.locator_data or {}

                try:
                    if field_def.field_type == FormFieldType.CHECKBOX.value:
                        # 复选框
                        checked = self._to_bool(value)
                        ok = self._fill_checkbox(doc, locator_type, locator_data, checked, field_def)
                    else:
                        # 文本/数值/日期/下拉/多行
                        text_val = str(value)
                        ok = self._fill_text_value(doc, locator_type, locator_data, text_val, field_def)
                    if ok:
                        filled_count += 1
                    else:
                        warnings.append(
                            f"字段「{field_def.field_label}」({field_key}) 定位失败，未填充"
                        )
                except Exception as e:
                    warnings.append(
                        f"字段「{field_def.field_label}」({field_key}) 填充异常: {e}"
                    )

            doc.save(output_path)

            # 4. 更新任务
            task.output_file_path = output_path
            task.output_filename = output_filename
            task.status = FormFillStatus.COMPLETED.value
            task.progress = 100
            task.completed_at = datetime.utcnow()
            task.updated_at = datetime.utcnow()
            db.commit()

            download_url = f"/api/form-fill/download/{task.id}"

            return {
                "task_id": task_id,
                "status": FormFillStatus.COMPLETED.value,
                "output_file": output_path,
                "output_filename": output_filename,
                "download_url": download_url,
                "filled_count": filled_count,
                "warnings": warnings,
                "message": f"模板填充完成，成功填充 {filled_count} 个字段",
            }
        finally:
            db.close()

    # ------------------------------------------------------------------
    # 通用单元格操作（从 iso_service 迁移/增强）
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell_text(cell, text: str, font_size: Optional[float] = None):
        """设置单元格文本，保留第一个 run 的格式，清除其余 run。

        相比 iso_service 版本的增强：
        - 保留第一个 run 的字体名称、大小、加粗、斜体、颜色
        - 支持多行文本（\\n 分段）
        - 可显式指定 font_size 覆盖
        """
        if not cell.paragraphs:
            cell.add_paragraph()

        # 取第一个段落作为锚点
        anchor_para = cell.paragraphs[0]

        # 深拷贝第一个 run 的格式（如果存在）
        ref_run_props = None
        if anchor_para.runs:
            ref_run_props = copy.deepcopy(anchor_para.runs[0]._element.rPr)

        # 清空段落里所有 run
        for run in anchor_para.runs:
            run.text = ""

        lines = text.split("\n") if text else [""]

        # 填充第一段（使用 anchor_para）
        FormFillEngine._write_line_to_paragraph(
            anchor_para, lines[0], ref_run_props, font_size
        )

        # 多余的段落清空内容
        for extra_para in cell.paragraphs[1:]:
            for run in extra_para.runs:
                run.text = ""

        # 如果有多行，在 anchor_para 后添加换行 run 或新段落
        # python-docx 中 add_break() 可在段内换行；若需新段落，用 add_paragraph
        for line in lines[1:]:
            # 段内换行（更贴合 Word 表格中常见的多行单元格）
            run = anchor_para.add_run()
            if ref_run_props is not None:
                run._element.rPr = copy.deepcopy(ref_run_props)
            if font_size:
                run.font.size = Pt(font_size)
            run.add_break()
            run.text = line

    @staticmethod
    def _write_line_to_paragraph(para, line_text: str, ref_run_props, font_size: Optional[float]):
        """向段落写入一行文本，优先复用已有 run 并保留格式。"""
        if para.runs:
            target_run = para.runs[0]
            target_run.text = line_text
            if font_size:
                target_run.font.size = Pt(font_size)
        else:
            new_run = para.add_run(line_text)
            if ref_run_props is not None:
                new_run._element.rPr = copy.deepcopy(ref_run_props)
            if font_size:
                new_run.font.size = Pt(font_size)
            elif ref_run_props is None:
                new_run.font.size = Pt(10.5)

    @staticmethod
    def _check_checkbox(cell, checked: bool):
        """勾选/取消复选框，将 □/☑/■ 替换为目标符号。

        勾选 -> ☑（也兼容用户要求的 ■ 风格：若原文使用 ■ 则保持 ■）
        取消 -> □
        """
        # 先判断原文使用哪种"已勾选"符号
        target_checked = "☑"
        for para in cell.paragraphs:
            for run in para.runs:
                if "■" in run.text:
                    target_checked = "■"
                    break

        target_unchecked = "□"
        mark = target_checked if checked else target_unchecked

        for para in cell.paragraphs:
            for run in para.runs:
                t = run.text
                if any(sym in t for sym in ("□", "☑", "■", "☐", "✓", "√")):
                    t = re.sub(r"[□☑■☐✓√]", mark, t)
                    run.text = t
                    return True
        return False

    @staticmethod
    def _set_bookmark_text(doc: Document, bookmark_name: str, text: str) -> bool:
        """在 Word 书签位置写入文本（替换书签所在 run 的内容）。"""
        from docx.oxml.ns import qn as _qn
        for bm in doc.element.body.iter(_qn("w:bookmarkStart")):
            if bm.get(_qn("w:name")) == bookmark_name:
                parent = bm.getparent()
                # 找到 bookmarkStart 之后的第一个 w:r（run），写入文本
                # 优先查找同一段落内的 run
                # 方法：从 bookmarkStart 向后遍历直到 bookmarkEnd
                bm_id = bm.get(_qn("w:id"))
                end_elem = None
                for sib in bm.itersiblings():
                    if sib.tag == _qn("w:bookmarkEnd") and sib.get(_qn("w:id")) == bm_id:
                        end_elem = sib
                        break
                    if sib.tag == _qn("w:r"):
                        # 清空该 run 内所有 t 节点，写入新文本
                        for t_elem in sib.findall(_qn("w:t")):
                            t_elem.text = text
                            # 保留空格
                            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                            return True
                        # run 内没有 w:t，则新建
                        from docx.oxml import OxmlElement
                        t_new = OxmlElement("w:t")
                        t_new.text = text
                        t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        sib.append(t_new)
                        return True
                # 没找到 run，则在 bookmarkStart 后面插入一个新 run
                from docx.oxml import OxmlElement
                new_r = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.text = text
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                new_r.append(new_t)
                bm.addnext(new_r)
                return True
        return False

    @staticmethod
    def _check_checkbox_by_bookmark(doc: Document, bookmark_name: str, checked: bool) -> bool:
        """在书签位置勾选复选框。"""
        from docx.oxml.ns import qn as _qn
        for bm in doc.element.body.iter(_qn("w:bookmarkStart")):
            if bm.get(_qn("w:name")) == bookmark_name:
                bm_id = bm.get(_qn("w:id"))
                # 在 bookmark 范围内查找 run 中的复选框字符
                target_checked = "☑"
                target_unchecked = "□"
                # 先扫描确定原风格
                for sib in bm.itersiblings():
                    if sib.tag == _qn("w:bookmarkEnd") and sib.get(_qn("w:id")) == bm_id:
                        break
                    if sib.tag == _qn("w:r"):
                        for t_elem in sib.findall(_qn("w:t")):
                            if t_elem.text and "■" in t_elem.text:
                                target_checked = "■"
                mark = target_checked if checked else target_unchecked
                for sib in bm.itersiblings():
                    if sib.tag == _qn("w:bookmarkEnd") and sib.get(_qn("w:id")) == bm_id:
                        break
                    if sib.tag == _qn("w:r"):
                        for t_elem in sib.findall(_qn("w:t")):
                            if t_elem.text and any(s in t_elem.text for s in ("□", "☑", "■", "☐", "✓", "√")):
                                t_elem.text = re.sub(r"[□☑■☐✓√]", mark, t_elem.text)
                                return True
                return False
        return False

    @staticmethod
    def _find_cell_by_label(
        doc: Document,
        label_text: str,
        direction: str = "right",
        table_idx: Optional[int] = None,
    ):
        """根据标签文本查找相邻单元格。

        :return: (cell, None) 或 (None, error_msg)
        """
        tables = [doc.tables[table_idx]] if table_idx is not None and table_idx < len(doc.tables) else doc.tables

        for table in tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if label_text.strip() in cell.text.replace("\n", "").replace(" ", ""):
                        # 找到标签所在单元格，返回 direction 方向的相邻单元格
                        try:
                            if direction == "right" and c_idx + 1 < len(row.cells):
                                return row.cells[c_idx + 1]
                            elif direction == "left" and c_idx - 1 >= 0:
                                return row.cells[c_idx - 1]
                            elif direction == "below" and r_idx + 1 < len(table.rows):
                                return table.rows[r_idx + 1].cells[c_idx]
                            elif direction == "above" and r_idx - 1 >= 0:
                                return table.rows[r_idx - 1].cells[c_idx]
                        except Exception:
                            continue
        return None

    @staticmethod
    def _find_checkbox_cell_by_label(
        doc: Document,
        label_text: str,
        direction: str = "right",
        table_idx: Optional[int] = None,
    ):
        """查找标签相邻的、含复选框符号的单元格。"""
        return FormFillEngine._find_cell_by_label(doc, label_text, direction, table_idx)

    # ------------------------------------------------------------------
    # 内部：按定位方式分发填充
    # ------------------------------------------------------------------

    def _fill_text_value(
        self,
        doc: Document,
        locator_type: str,
        locator_data: Dict[str, Any],
        text: str,
        field_def: FormField,
    ) -> bool:
        """填充普通文本字段。"""
        lt = locator_type

        if lt == LocatorType.TABLE_CELL.value:
            table_idx = int(locator_data.get("table_idx", 0))
            row = int(locator_data.get("row", 0))
            col = int(locator_data.get("col", 0))
            if table_idx >= len(doc.tables):
                return False
            table = doc.tables[table_idx]
            if row >= len(table.rows):
                return False
            r = table.rows[row]
            if col >= len(r.cells):
                return False
            cell = r.cells[col]
            self._set_cell_text(cell, text)
            return True

        elif lt == LocatorType.BOOKMARK.value:
            bookmark = locator_data.get("bookmark")
            if not bookmark:
                return False
            return self._set_bookmark_text(doc, bookmark, text)

        elif lt == LocatorType.LABEL_ADJACENT.value:
            label_text = locator_data.get("label_text", "")
            direction = locator_data.get("direction", "right")
            table_idx = locator_data.get("table_idx")
            cell = self._find_cell_by_label(doc, label_text, direction, table_idx)
            if cell is None:
                return False
            self._set_cell_text(cell, text)
            return True

        return False

    def _fill_checkbox(
        self,
        doc: Document,
        locator_type: str,
        locator_data: Dict[str, Any],
        checked: bool,
        field_def: FormField,
    ) -> bool:
        """填充复选框字段。"""
        lt = locator_type

        if lt == LocatorType.TABLE_CELL.value:
            table_idx = int(locator_data.get("table_idx", 0))
            row = int(locator_data.get("row", 0))
            col = int(locator_data.get("col", 0))
            if table_idx >= len(doc.tables):
                return False
            table = doc.tables[table_idx]
            if row >= len(table.rows):
                return False
            r = table.rows[row]
            if col >= len(r.cells):
                return False
            cell = r.cells[col]
            ok = self._check_checkbox(cell, checked)
            if not ok:
                # 单元格里没有现成复选框符号，则在文本前追加
                mark = "■" if checked else "□"
                existing = cell.text.strip()
                self._set_cell_text(cell, f"{mark} {existing}".strip())
                return True
            return True

        elif lt == LocatorType.BOOKMARK.value:
            bookmark = locator_data.get("bookmark")
            if not bookmark:
                return False
            return self._check_checkbox_by_bookmark(doc, bookmark, checked)

        elif lt == LocatorType.LABEL_ADJACENT.value:
            label_text = locator_data.get("label_text", "")
            direction = locator_data.get("direction", "right")
            table_idx = locator_data.get("table_idx")
            cell = self._find_checkbox_cell_by_label(doc, label_text, direction, table_idx)
            if cell is None:
                return False
            ok = self._check_checkbox(cell, checked)
            if not ok:
                mark = "■" if checked else "□"
                existing = cell.text.strip()
                self._set_cell_text(cell, f"{mark} {existing}".strip())
                return True
            return True

        return False

    # ------------------------------------------------------------------
    # 辅助方法
    # ------------------------------------------------------------------

    @staticmethod
    def _to_bool(val: Any) -> bool:
        """将多种表达转换为 bool。"""
        if isinstance(val, bool):
            return val
        if val is None:
            return False
        s = str(val).strip().lower()
        return s in ("true", "1", "yes", "y", "是", "勾选", "√", "☑", "■", "on")

    def get_task(self, task_id: int) -> Optional[FormFillTask]:
        db = SessionLocal()
        try:
            return db.query(FormFillTask).filter(FormFillTask.id == task_id).first()
        finally:
            db.close()

    def get_output_file(self, task_id: int) -> Optional[str]:
        task = self.get_task(task_id)
        if task and task.output_file_path and os.path.exists(task.output_file_path):
            return task.output_file_path
        return None

    def list_task_field_values(self, task_id: int) -> List[Dict[str, Any]]:
        """获取任务的所有字段值（用于前端展示/确认）。"""
        db = SessionLocal()
        try:
            task = db.query(FormFillTask).filter(FormFillTask.id == task_id).first()
            if not task:
                return []
            fields = db.query(FormField).filter(
                FormField.template_id == task.template_id
            ).order_by(FormField.sort_order, FormField.id).all()
            result = []
            for f in fields:
                fv = db.query(FormFillFieldValue).filter(
                    FormFillFieldValue.task_id == task_id,
                    FormFillFieldValue.field_id == f.id,
                ).first()
                result.append({
                    "field_id": f.id,
                    "field_key": f.field_key,
                    "field_label": f.field_label,
                    "field_type": f.field_type,
                    "required": f.required,
                    "extracted_value": fv.extracted_value if fv else None,
                    "confirmed_value": fv.confirmed_value if fv else None,
                    "confidence": fv.confidence if fv else None,
                    "source_text": fv.source_text if fv else None,
                    "source_section": fv.source_section if fv else None,
                })
            return result
        finally:
            db.close()


# ------------------------------------------------------------------
# 单例
# ------------------------------------------------------------------

form_fill_engine = FormFillEngine()
