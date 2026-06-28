import os
import re
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from app.core.config import settings


@dataclass
class ParsedTable:
    """解析后的表格"""
    table_index: int
    page_number: int = 0
    section_path: str = ""
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    caption: str = ""


@dataclass
class ParsedChapter:
    """解析后的章节"""
    level: int
    number: str
    title: str
    page_start: int = 0
    content_text: str = ""
    tables: List[ParsedTable] = field(default_factory=list)
    children: List["ParsedChapter"] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    file_path: str
    file_type: str
    total_pages: int = 0
    full_text: str = ""
    pages: List[Dict] = field(default_factory=list)
    chapters: List[ParsedChapter] = field(default_factory=list)
    tables: List[ParsedTable] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)


class DocumentParser:
    """增强版文档解析器，支持页码、章节、表格结构化、Excel"""
    
    def __init__(self):
        # 修复5：默认值与config.py保持一致（500/100）
        self.chunk_size = getattr(settings, 'CHUNK_SIZE', 500)
        self.chunk_overlap = getattr(settings, 'CHUNK_OVERLAP', 100)
        # 章节标题正则模式
        self.chapter_patterns = [
            re.compile(r'^第[一二三四五六七八九十百]+[章节篇编]\s+(.+)$'),
            re.compile(r'^(\d+(?:\.\d+)*)\s+([^\d\.].+)$'),
            re.compile(r'^(\d+(?:\.\d+)*)\s*[、\.]\s*(.+)$'),
            re.compile(r'^[（(]?[一二三四五六七八九十]+[)）]\s*[、\.]?\s*(.+)$'),
        ]
    
    def parse(self, file_path: str, file_type: str = None) -> str:
        """解析文档，返回全文本（兼容旧接口）"""
        result = self.parse_structured(file_path)
        return result.full_text
    
    def parse_structured(self, file_path: str) -> ParsedDocument:
        """解析文档，返回结构化结果"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf_structured(file_path)
        elif ext == '.docx':
            return self._parse_docx_structured(file_path)
        elif ext == '.doc':
            # 修复5：python-docx不支持.doc格式，返回明确提示
            doc = ParsedDocument(file_path=file_path, file_type="doc")
            doc.full_text = "[不支持的格式] .doc 是旧版Word格式，请转换为.docx后再上传。"
            return doc
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel_structured(file_path)
        elif ext == '.txt':
            return self._parse_txt_structured(file_path)
        else:
            return self._parse_txt_structured(file_path)
    
    def parse_and_chunk(self, file_path: str, document_id: int, 
                        file_name: str = "") -> List[Dict[str, Any]]:
        """解析文档并分块（兼容旧接口，使用新的章节感知分块）"""
        parsed = self.parse_structured(file_path)
        chunks = self._chapter_aware_chunk(parsed)
        
        result = []
        for i, chunk in enumerate(chunks):
            result.append({
                "document_id": document_id,
                "chunk_text": chunk["text"],
                "page_number": chunk.get("page_number"),
                "section_title": chunk.get("section_title", file_name),
                "chapter_path": chunk.get("chapter_path"),
                "tables_json": chunk.get("tables_json"),
                "chunk_index": i,
            })
        return result
    
    def parse_and_chunk_enhanced(self, file_path: str) -> tuple:
        """增强版解析和分块，返回 (ParsedDocument, chunks)"""
        parsed = self.parse_structured(file_path)
        chunks = self._chapter_aware_chunk(parsed)
        return parsed, chunks
    
    def _parse_pdf_structured(self, file_path: str) -> ParsedDocument:
        """结构化解析PDF - 修复5+6：优先fitz(PyMuPDF)，其次pdfplumber，最后PyPDF2，支持表格提取"""
        doc = ParsedDocument(file_path=file_path, file_type="pdf")
        
        # 第一优先级：fitz (PyMuPDF) - 性能最佳
        try:
            import fitz
            pdf_doc = fitz.open(file_path)
            doc.total_pages = len(pdf_doc)
            all_text_parts = []
            current_page_texts = []
            
            for page_num, page in enumerate(pdf_doc, 1):
                page_text = page.get_text()
                current_page_texts.append({"page_number": page_num, "text": page_text})
                all_text_parts.append(page_text)
            
            doc.pages = current_page_texts
            doc.full_text = "\n".join(all_text_parts)
            doc.chapters = self._extract_chapters(doc.full_text)
            pdf_doc.close()
            return doc
        except ImportError:
            print("[PDF解析] PyMuPDF(fitz)未安装，尝试pdfplumber...")
        except Exception as e:
            print(f"fitz解析PDF失败: {e}，尝试pdfplumber...")
        
        # 第二优先级：pdfplumber（支持表格提取）
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                doc.total_pages = len(pdf.pages)
                all_text_parts = []
                current_page_texts = []
                table_index = 0
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    current_page_texts.append({"page_number": page_num, "text": page_text})
                    all_text_parts.append(page_text)
                    
                    # 使用pdfplumber提取表格
                    try:
                        tables_on_page = page.extract_tables() or []
                        for tbl in tables_on_page:
                            if not tbl or len(tbl) < 2:
                                continue
                            parsed_table = ParsedTable(
                                table_index=table_index,
                                page_number=page_num,
                            )
                            parsed_table.headers = [str(cell).strip() if cell else "" for cell in tbl[0]]
                            parsed_table.rows = [
                                [str(cell).strip() if cell else "" for cell in row]
                                for row in tbl[1:]
                            ]
                            parsed_table.rows = [r for r in parsed_table.rows if any(c for c in r)]
                            if parsed_table.rows:
                                doc.tables.append(parsed_table)
                                table_text = self._table_to_text(parsed_table)
                                if table_text:
                                    all_text_parts.append(table_text)
                                table_index += 1
                    except Exception as te:
                        print(f"PDF表格提取失败(页{page_num}): {te}")
                
                doc.pages = current_page_texts
                doc.full_text = "\n".join(all_text_parts)
                doc.chapters = self._extract_chapters(doc.full_text)
            return doc
        except ImportError:
            print("[PDF解析] pdfplumber未安装，尝试PyPDF2...")
        except Exception as e:
            print(f"pdfplumber解析PDF失败: {e}，尝试PyPDF2...")
        
        # 第三优先级：PyPDF2（纯文本，无表格）
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            doc.total_pages = len(reader.pages)
            all_text_parts = []
            current_page_texts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                current_page_texts.append({"page_number": page_num, "text": page_text})
                all_text_parts.append(page_text)
            
            doc.pages = current_page_texts
            doc.full_text = "\n".join(all_text_parts)
            doc.chapters = self._extract_chapters(doc.full_text)
            return doc
        except Exception as e:
            print(f"PyPDF2解析PDF也失败: {e}")
            doc.full_text = ""
            return doc
    
    def _parse_docx_structured(self, file_path: str) -> ParsedDocument:
        """结构化解析DOCX"""
        doc = ParsedDocument(file_path=file_path, file_type="docx")
        try:
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(file_path)
            
            all_text_parts = []
            chapters = []
            tables = []
            table_index = 0
            current_chapter = None
            current_section_path = ""
            para_index = 0
            total_paras = len(docx_doc.paragraphs)
            
            # 先估算总页数（近似值）
            doc.total_pages = max(1, total_paras // 30)
            
            # 解析段落
            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 检测章节标题
                chapter_info = self._detect_chapter_heading(text, para.style.name if para.style else "")
                if chapter_info:
                    level, number, title = chapter_info
                    new_chapter = ParsedChapter(
                        level=level, number=number, title=title,
                        page_start=max(1, para_index * doc.total_pages // max(total_paras, 1))
                    )
                    chapters.append(new_chapter)
                    current_chapter = new_chapter
                    current_section_path = f"{number} {title}"
                
                all_text_parts.append(text)
                if current_chapter:
                    if current_chapter.content_text:
                        current_chapter.content_text += "\n"
                    current_chapter.content_text += text
                
                para_index += 1
            
            # 解析表格
            for t_idx, table in enumerate(docx_doc.tables):
                parsed_table = ParsedTable(
                    table_index=table_index,
                    page_number=0,
                    section_path=current_section_path
                )
                
                rows_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_cells)
                
                if rows_data:
                    parsed_table.headers = rows_data[0]
                    parsed_table.rows = rows_data[1:]
                    tables.append(parsed_table)
                    
                    # 尝试向前查找表题
                    table_caption = ""
                    for prev_text in reversed(all_text_parts[-20:]):
                        if re.search(r'表\s*\d+[\.\-]\d*', prev_text):
                            table_caption = prev_text
                            break
                    parsed_table.caption = table_caption
                    
                    if current_chapter:
                        current_chapter.tables.append(parsed_table)
                
                table_index += 1
            
            doc.full_text = "\n".join(all_text_parts)
            doc.chapters = chapters
            doc.tables = tables
            
        except Exception as e:
            print(f"DOCX结构化解析失败: {e}")
            doc.full_text = self._parse_docx(file_path)
        
        return doc
    
    def _parse_excel_structured(self, file_path: str) -> ParsedDocument:
        """结构化解析Excel"""
        doc = ParsedDocument(file_path=file_path, file_type="excel")
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            all_text_parts = []
            tables = []
            table_index = 0
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parsed_table = ParsedTable(
                    table_index=table_index,
                    page_number=1,
                    section_path=sheet_name,
                    caption=sheet_name
                )
                
                rows_data = []
                for row in ws.iter_rows(values_only=True):
                    row_cells = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_cells):
                        rows_data.append(row_cells)
                
                if rows_data:
                    # 尝试识别表头
                    parsed_table.headers = rows_data[0] if rows_data else []
                    parsed_table.rows = rows_data[1:] if len(rows_data) > 1 else []
                    tables.append(parsed_table)
                    
                    # 添加到全文
                    all_text_parts.append(f"=== {sheet_name} ===")
                    for row in rows_data:
                        all_text_parts.append(" | ".join(row))
                
                table_index += 1
            
            doc.total_pages = len(wb.sheetnames)
            doc.full_text = "\n".join(all_text_parts)
            doc.tables = tables
            
            wb.close()
        except Exception as e:
            print(f"Excel解析失败: {e}")
            doc.full_text = ""
        
        return doc
    
    def _parse_txt_structured(self, file_path: str) -> ParsedDocument:
        """结构化解析TXT"""
        doc = ParsedDocument(file_path=file_path, file_type="txt")
        text = self._parse_txt(file_path)
        doc.full_text = text
        doc.total_pages = 1
        doc.pages = [{"page_number": 1, "text": text}]
        doc.chapters = self._extract_chapters(text)
        return doc
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件（兼容旧接口）- 修复5+6：优先fitz，其次pdfplumber，最后PyPDF2"""
        # 第一优先级：fitz (PyMuPDF)
        try:
            import fitz
            pdf_doc = fitz.open(file_path)
            text_parts = []
            for page in pdf_doc:
                text_parts.append(page.get_text())
            pdf_doc.close()
            return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception as e:
            print(f"fitz解析PDF失败: {e}")
        
        # 第二优先级：pdfplumber（含表格）
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    for tbl in (page.extract_tables() or []):
                        for row in tbl:
                            row_text = " | ".join(str(c).strip() if c else "" for c in row)
                            if row_text.strip(" |"):
                                text_parts.append(row_text)
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception as e:
            print(f"pdfplumber解析PDF失败: {e}")
        
        # 第三优先级：PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"PDF解析失败: {e}")
            return ""
    
    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档（兼容旧接口）"""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"Word解析失败: {e}")
            return ""
    
    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                print(f"文本解析失败: {e}")
                return ""
    
    def _detect_chapter_heading(self, text: str, style_name: str = "") -> Optional[tuple]:
        """检测是否为章节标题，返回 (level, number, title) 或 None"""
        # 优先通过样式判断
        if style_name:
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.replace('Heading', '').strip())
                    if level >= 1 and level <= 4:
                        # 尝试提取编号
                        match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
                        if match:
                            return (level, match.group(1), match.group(2))
                        return (level, "", text)
                except ValueError:
                    pass
        
        # 通过正则匹配
        for pattern in self.chapter_patterns:
            match = pattern.match(text.strip())
            if match:
                groups = match.groups()
                if len(groups) == 1:
                    # "第一章 xxx" 格式
                    return (1, "", groups[0])
                elif len(groups) >= 2:
                    number = groups[0]
                    title = groups[1] if len(groups) > 1 else ""
                    # 根据编号点号数量判断层级
                    level = number.count('.') + 1 if number and re.match(r'^\d', number) else 1
                    level = min(level, 4)
                    return (level, number, title)
        
        return None
    
    def _extract_chapters(self, text: str) -> List[ParsedChapter]:
        """从纯文本中提取章节结构"""
        chapters = []
        lines = text.split('\n')
        current_chapter = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            chapter_info = self._detect_chapter_heading(line)
            if chapter_info:
                level, number, title = chapter_info
                new_chapter = ParsedChapter(level=level, number=number, title=title)
                chapters.append(new_chapter)
                current_chapter = new_chapter
            elif current_chapter:
                if current_chapter.content_text:
                    current_chapter.content_text += "\n"
                current_chapter.content_text += line
        
        return chapters
    
    def _extract_tables_from_text(self, text: str, start_index: int = 0) -> List[ParsedTable]:
        """从文本中简单提取表格（按连续|分隔行判断）- 主要用于非PDF格式"""
        tables = []
        # 修复5：PDF表格已在_parse_pdf_structured中通过pdfplumber提取，此处仅处理纯文本中的表格
        lines = text.split('\n')
        table_lines = []
        for line in lines:
            if '|' in line and line.count('|') >= 2:
                table_lines.append(line)
            else:
                if len(table_lines) >= 2:
                    # 尝试解析为表格
                    rows = []
                    for tl in table_lines:
                        cells = [c.strip() for c in tl.split('|') if c.strip()]
                        if cells:
                            rows.append(cells)
                    if len(rows) >= 2:
                        t = ParsedTable(table_index=start_index + len(tables))
                        t.headers = rows[0]
                        t.rows = rows[1:]
                        if any(any(c for c in r) for r in t.rows):
                            tables.append(t)
                table_lines = []
        return tables
    
    @staticmethod
    def _table_to_text(table: ParsedTable) -> str:
        """将ParsedTable转换为文本表示"""
        lines = []
        if table.caption:
            lines.append(table.caption)
        if table.headers:
            lines.append(" | ".join(table.headers))
        for row in table.rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)
    
    def _chapter_aware_chunk(self, parsed: ParsedDocument) -> List[Dict[str, Any]]:
        """章节感知分块：先按章节切分，章节内滑动窗口"""
        chunks = []
        
        if parsed.chapters:
            # 有章节结构，按章节分块
            for chapter in parsed.chapters:
                chapter_title = f"{chapter.number} {chapter.title}".strip()
                chapter_chunks = self._split_text_into_chunks(
                    chapter.content_text,
                    self.chunk_size,
                    self.chunk_overlap
                )
                
                # 章节内的表格
                tables_json = [
                    {"headers": t.headers, "rows": t.rows, "caption": t.caption}
                    for t in chapter.tables
                ] if chapter.tables else None
                
                for chunk_text in chapter_chunks:
                    chunks.append({
                        "text": chunk_text,
                        "page_number": chapter.page_start,
                        "section_title": chapter_title,
                        "chapter_path": chapter.number,
                        "tables_json": tables_json
                    })
                
                # 处理子章节（递归）
                for child in chapter.children:
                    child_chunks = self._chapter_aware_chunk_chapter(child)
                    chunks.extend(child_chunks)
        else:
            # 无章节结构，使用简单分块
            simple_chunks = self._chunk_text(parsed.full_text)
            for chunk_text in simple_chunks:
                chunks.append({
                    "text": chunk_text,
                    "page_number": None,
                    "section_title": "",
                    "chapter_path": None,
                    "tables_json": None
                })
        
        return chunks
    
    def _chapter_aware_chunk_chapter(self, chapter: ParsedChapter, parent_path: str = "") -> List[Dict[str, Any]]:
        """递归处理子章节分块"""
        chunks = []
        full_path = f"{parent_path} > {chapter.number} {chapter.title}".strip(" >")
        
        chapter_chunks = self._split_text_into_chunks(
            chapter.content_text,
            self.chunk_size,
            self.chunk_overlap
        )
        
        for chunk_text in chapter_chunks:
            chunks.append({
                "text": chunk_text,
                "page_number": chapter.page_start,
                "section_title": full_path,
                "chapter_path": chapter.number,
                "tables_json": None
            })
        
        for child in chapter.children:
            chunks.extend(self._chapter_aware_chunk_chapter(child, full_path))
        
        return chunks
    
    def _split_text_into_chunks(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """将文本按固定大小滑动窗口分块"""
        if not text:
            return []
        
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += "\n"
                current_chunk += para
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                if len(para) > chunk_size:
                    # 修复：超长段落切分后需要正确设置current_chunk为overlap尾部，避免丢失
                    start = 0
                    last_tail = ""
                    while start < len(para):
                        end = start + chunk_size
                        chunk = para[start:end]
                        if chunk:
                            chunks.append(chunk)
                            # 保存尾部用于overlap
                            if len(chunk) >= overlap:
                                last_tail = chunk[-overlap:]
                            else:
                                last_tail = chunk
                        start = end - overlap
                    # 将overlap尾部设为current_chunk，使后续段落能正确衔接
                    current_chunk = last_tail if last_tail else ""
                else:
                    current_chunk = para
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _chunk_text(self, text: str) -> List[str]:
        """文本分块（兼容旧接口）"""
        return self._split_text_into_chunks(text, self.chunk_size, self.chunk_overlap)


# 全局单例
document_parser = DocumentParser()
